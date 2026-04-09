from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

C_DARK    = RGBColor(0x1E, 0x29, 0x3B)
C_MID     = RGBColor(0x33, 0x41, 0x55)
C_ACCENT  = RGBColor(0x7C, 0x3A, 0xED)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_BODY    = RGBColor(0x1E, 0x29, 0x3B)
C_MUTED   = RGBColor(0x47, 0x55, 0x69)
C_HIGH    = RGBColor(0xFE, 0xE2, 0xE2)
C_MED_C   = RGBColor(0xFE, 0xF9, 0xC3)
C_LOW     = RGBColor(0xDC, 0xFC, 0xE7)
C_HDR_BG  = RGBColor(0x1E, 0x29, 0x3B)
C_SEC_BG  = RGBColor(0x33, 0x41, 0x55)
C_ALT     = RGBColor(0xF8, 0xFA, 0xFC)
C_TCID    = RGBColor(0xEF, 0xF6, 0xFF)

def hex_to_rgb_str(c): return '{:02X}{:02X}{:02X}'.format(c[0], c[1], c[2])

def shade_cell(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_to_rgb_str(color))
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_cell_border(cell, color="CBD5E1", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = C_DARK
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = C_MUTED
    p.paragraph_format.space_after = Pt(16)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '7C3AED')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_ACCENT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_MID
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10)
    run.font.color.rgb = color or C_BODY
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = C_BODY
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_MUTED
    p.paragraph_format.space_after = Pt(6)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(10)

def make_table(doc, headers, rows, col_widths, header_bg=None, font_size=9):
    if header_bg is None:
        header_bg = C_HDR_BG
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hrow.cells[i]
        cell.width = Inches(w)
        shade_cell(cell, header_bg)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = table.add_row()
        bg = C_WHITE if ri % 2 == 0 else C_ALT
        for ci, (val, w) in enumerate(zip(row, col_widths)):
            cell = tr.cells[ci]
            cell.width = Inches(w)
            shade_cell(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.color.rgb = C_BODY
            if ci == 0:
                run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_title(doc, "VABERP")
add_subtitle(doc, "QA Handover Document & Test Cases")
add_subtitle(doc, "Version 2.0  |  2026-04-07  |  185 Test Cases  |  Prepared by: Senior QA Lead")
add_hr(doc)

add_h1(doc, "PART 1 — QA HANDOVER DOCUMENT")

# ── Section 1: Project Overview ───────────────────────────────────────────────
add_h2(doc, "1. Project Overview")
make_table(doc,
    ["Field", "Details"],
    [
        ("Product Name", "VABERP — VAB Enterprise ERP"),
        ("Platform", "Web-based (Next.js 14 Frontend · NestJS Backend · PostgreSQL)"),
        ("App URL", "http://64.235.43.187:2262"),
        ("API URL", "http://64.235.43.187:2261"),
        ("Super Admin Portal", "http://64.235.43.187:2262/superadmin/login"),
        ("Reseller Portal", "http://64.235.43.187:2262/reseller/login"),
        ("Enterprise Portal", "http://64.235.43.187:2262/login"),
        ("Purpose", "End-to-end manufacturing operations management with multi-tenant SaaS architecture"),
        ("Primary Users", "Super Admin, Reseller Admin, Enterprise Admin, Manager, Employee"),
        ("Core Workflow", "RFQ → Quotation → Approval → Purchase Order → Invoice"),
        ("Current Status", "Active / In QA Handover"),
    ],
    [1.8, 4.7]
)
add_body(doc, "VABERP is a multi-tenant manufacturing ERP system. The platform has three separate portals: a Super Admin portal for platform-level management, a Reseller Admin portal for managing tenant enterprises, and an Enterprise portal for day-to-day manufacturing operations.")

# ── Section 2: User Roles ─────────────────────────────────────────────────────
add_h2(doc, "2. User Roles & Permissions")
add_h3(doc, "Role Hierarchy (5 Tiers)")
add_body(doc, "Super Admin  →  Reseller Admin  →  Enterprise Admin  →  Manager  →  Employee", bold=True)
add_note(doc, "Super Admin manages the entire platform. Reseller Admin manages a portfolio of enterprise tenants. Enterprise Admin, Manager, and Employee operate within a single enterprise.")

add_h3(doc, "Portal & Guard Architecture")
make_table(doc,
    ["Role", "Login Portal", "JWT Guard", "Scope"],
    [
        ("Super Admin", "/superadmin/login", "SuperAdminGuard (type=super_admin)", "Full platform — all enterprises, all resellers"),
        ("Reseller Admin", "/reseller/login", "ResellerGuard (type=reseller)", "Own tenant portfolio only"),
        ("Enterprise Admin", "/login", "JwtAuthGuard (role=admin)", "Full enterprise — all modules, all users"),
        ("Manager", "/login", "JwtAuthGuard (role=manager)", "Own records + direct report employees"),
        ("Employee", "/login", "JwtAuthGuard (role=employee)", "Own assigned records only"),
    ],
    [1.2, 1.6, 2.4, 2.4]
)

add_h3(doc, "Permission Matrix — Enterprise Level (Admin / Manager / Employee)")
make_table(doc,
    ["Permission", "Admin", "Manager", "Employee"],
    [
        ("View ALL records (all users)", "YES", "NO", "NO"),
        ("View own team's records", "YES", "YES (own team)", "NO"),
        ("View own assigned records", "YES", "YES", "YES"),
        ("Create RFQ / Quotation", "YES", "YES", "YES"),
        ("Submit Quotation for Approval", "YES", "YES", "YES"),
        ("Approve / Reject Quotation", "YES", "YES", "NO"),
        ("Create Purchase Order", "YES", "YES", "NO"),
        ("Generate Invoice", "YES", "YES", "NO"),
        ("Record Payments", "YES", "YES", "NO"),
        ("Manage Users / Roles", "YES", "NO", "NO"),
        ("Access System Settings", "YES", "Limited", "NO"),
        ("View Reports", "YES", "YES (team only)", "NO"),
    ],
    [3.0, 0.9, 1.1, 1.1]
)

add_h3(doc, "Super Admin Capabilities")
for rule in [
    "Full platform access — views and manages ALL enterprises and ALL resellers.",
    "Creates and manages reseller accounts, subscription plans, and coupon codes.",
    "Approves, locks, or unlocks enterprise accounts.",
    "Credits reseller wallets and manages platform-level financial accounts.",
    "Views all support tickets from all enterprises and resellers.",
    "Manages the services/features catalog available to subscription plans.",
    "CANNOT access enterprise ERP modules (quotations, POs, invoices etc.) — portal is fully separate.",
]:
    add_bullet(doc, rule)

add_h3(doc, "Reseller Admin Capabilities")
for rule in [
    "Manages only the enterprises (tenants) assigned to them by Super Admin.",
    "Creates new tenant enterprises under their portfolio.",
    "Assigns subscription plans to tenants and manages renewals.",
    "Views wallet balance, commission earnings, and billing history.",
    "CANNOT see other resellers' tenants — strict data isolation.",
    "CANNOT access Super Admin portal or enterprise ERP modules.",
    "Subject to a maximum tenant limit defined by their own reseller plan.",
]:
    add_bullet(doc, rule)

add_h3(doc, "Enterprise Role Rules (Admin / Manager / Employee)")
for rule in [
    "Admin has unrestricted visibility and action access across all ERP modules.",
    "Manager sees their own records AND all records of their direct report employees. Cannot see other managers' team data.",
    "Employee sees only records directly assigned to them.",
    "Role assignment is performed exclusively by Enterprise Admin.",
    "Hierarchy changes (reassigning employees) take effect immediately.",
]:
    add_bullet(doc, rule)

# ── Section 3: Module Breakdown ───────────────────────────────────────────────
add_h2(doc, "3. Module Breakdown")

add_h3(doc, "Super Admin Modules")
sa_modules = [
    ("Dashboard", "Platform KPIs: Total Enterprises, Active Enterprises, Total Resellers, Revenue, Monthly Charts."),
    ("Enterprises", "Create, view, approve, lock/unlock all enterprise accounts. Reassign enterprise to different reseller. Manage subscription expiry."),
    ("Resellers", "Create and manage reseller accounts. Lock/unlock reseller accounts. Manage reseller plans and wallets."),
    ("Subscriptions", "Create and manage subscription plans with feature sets, limits, and pricing."),
    ("Accounts", "Platform-level financial summary with period filtering."),
    ("Services", "Manage the services/features catalog available across subscription plans."),
    ("Coupons", "Create and manage discount coupon codes for enterprise subscriptions."),
    ("Support", "View and reply to all support tickets from all enterprises and resellers."),
    ("Employees (Platform)", "View all employees across all enterprises."),
]
for title, desc in sa_modules:
    add_body(doc, f"  {title}: {desc}")

add_h3(doc, "Reseller Admin Modules")
ra_modules = [
    ("Dashboard", "KPIs: Total Tenants, Active/Expired Subscriptions, Revenue, Commission Earned, Wallet Balance."),
    ("Tenants", "View, create, and manage own tenant enterprises. Assign plans and manage subscriptions."),
    ("Subscriptions", "View subscription status of all tenants. Renew expiring subscriptions."),
    ("Plans", "View available plans, set custom pricing per plan."),
    ("Wallet", "View wallet balance and full transaction history."),
    ("Commissions", "View commission earnings by tenant, plan, and date."),
    ("Usage", "View employee counts and usage stats per tenant."),
    ("Billing", "View billing/invoice history."),
    ("Reports", "Revenue, commission, and tenant activity reports."),
    ("Profile", "Update own reseller profile details."),
]
for title, desc in ra_modules:
    add_body(doc, f"  {title}: {desc}")

add_h3(doc, "Enterprise ERP Modules")
erp_modules = [
    ("3.1 RFQ — Request for Quotation", "Allows procurement team to request pricing from vendors. An RFQ can be sent to multiple vendors. Once a vendor responds, the RFQ is converted into a formal Quotation.\nKey states: Draft → Sent → Received → Converted"),
    ("3.2 Quotation", "Core commercial document built with line items (product, qty, unit price, discount). Supports full version control — version increments ONLY on user-initiated revision, never on rejection.\nKey states: Draft → Submitted → Approved | Rejected → Revised (new version)"),
    ("3.3 Purchase Order (PO)", "Generated from an Approved Quotation. Represents formal vendor commitment. Visibility strictly controlled — employees see only POs assigned to them.\nKey states: Draft → Ordered → Received → Invoiced"),
    ("3.4 Invoice", "Generated from a completed PO or Sales Order. Supports partial and full payment recording. Once fully paid, invoice becomes read-only.\nKey states: Unpaid → Partially Paid → Paid"),
    ("3.5 Sales Module", "Manages customer-facing orders from enquiry through production to delivery. Linked to manufacturing job cards and dispatch records.\nKey states: New → In Progress → Dispatched → Completed | Cancelled"),
    ("3.6 Employee & Role Management", "Admin creates and manages user accounts, assigns roles, and sets the reporting structure. Changes to hierarchy are reflected immediately across all modules."),
]
for title, desc in erp_modules:
    add_h3(doc, title)
    for line in desc.split('\n'):
        if line.strip():
            add_body(doc, line)

# ── Section 4: Workflows ──────────────────────────────────────────────────────
add_h2(doc, "4. End-to-End Workflows")

add_h3(doc, "Workflow 1: Procurement Flow — RFQ → PO → Invoice")
for i, s in enumerate([
    "User raises RFQ → selects vendor, adds line items, submits",
    "Vendor responds → RFQ status = Received",
    "User converts RFQ to Quotation → Quotation v1 created, items pre-filled",
    "User submits Quotation for approval → status = Submitted",
    "[Manager/Admin rejects] → Status = Rejected, VERSION STAYS SAME",
    "[User revises] → New version created (v1 → v2), resubmitted",
    "Manager/Admin approves → Status = Approved",
    "User creates Purchase Order from Approved Quotation",
    "Goods received → PO status updated to Received",
    "Invoice generated from PO → status = Unpaid",
    "Payment recorded → Partially Paid → Paid",
], 1):
    add_bullet(doc, f"Step {i}: {s}")

add_h3(doc, "Workflow 2: Sales Flow — Enquiry → Invoice")
for i, s in enumerate([
    "Customer enquiry captured in system",
    "Quotation created for customer with line items",
    "Quotation approved by Manager/Admin",
    "Sales Order created from Approved Quotation",
    "Job Card generated for manufacturing team",
    "Production stages completed",
    "Goods dispatched → dispatch record created",
    "Invoice raised to customer",
    "Payment received → Invoice marked Paid",
], 1):
    add_bullet(doc, f"Step {i}: {s}")

add_h3(doc, "Workflow 3: Quotation Version Control")
add_body(doc, "Create Quotation → v1 (Draft) → Submit → v1 (Submitted)")
add_bullet(doc, "Approved → v1 (Approved) — version stays 1")
add_bullet(doc, "Rejected → v1 (Rejected) — version STAYS 1, NOT incremented")
add_bullet(doc, "User Revises → v2 (Draft) → Submit → v2 (Submitted) → Approved/Rejected")
add_note(doc, "★ CRITICAL: Version increments ONLY when a user submits a revision — NEVER on rejection alone.")

add_h3(doc, "Workflow 4: Super Admin → Reseller → Enterprise Setup")
for i, s in enumerate([
    "Super Admin creates a subscription plan (features, limits, price)",
    "Super Admin creates a Reseller account",
    "Reseller logs in at /reseller/login",
    "Reseller creates a new Tenant (Enterprise) and assigns a plan",
    "Enterprise Admin logs in at /login",
    "Enterprise Admin creates Manager and Employee accounts",
    "Enterprise operations begin (RFQ, Quotations, POs, Invoices)",
], 1):
    add_bullet(doc, f"Step {i}: {s}")

# ── Section 5: Feature Expectations ──────────────────────────────────────────
add_h2(doc, "5. Feature Expectations")
make_table(doc,
    ["Feature", "Expected System Behavior"],
    [
        ("Quotation Version Control", "Version increments ONLY on user revision. Rejection by approver NEVER increments version."),
        ("RFQ to Quotation Conversion", "Converts to Quotation v1 with all line items, quantities, vendor details pre-filled."),
        ("PO Creation Gate", "PO can only be created from Approved quotations. Draft/Submitted/Rejected block PO creation."),
        ("PO Visibility — Employee", "Employee sees ONLY POs where they are the explicitly assigned owner."),
        ("PO Visibility — Manager", "Manager sees all POs belonging to themselves and their direct report employees."),
        ("Invoice Auto-Population", "Invoice amount is pre-populated from PO total. No manual entry required."),
        ("Invoice Status Flow", "Unpaid → Partially Paid → Paid. No reverse transitions allowed."),
        ("Paid Invoice Lock", "Once Paid, all invoice fields are read-only. No further payment edits."),
        ("Duplicate PO Prevention", "Only one PO per approved Quotation. Second attempt must be blocked or warned."),
        ("Duplicate Invoice Prevention", "Only one Invoice per PO. Second attempt must be blocked."),
        ("RBAC Enforcement", "Restricted pages/API endpoints must return 403 or redirect. No data exposed via URL manipulation."),
        ("Manager Cross-Visibility Block", "Manager must never see another manager's team records, even by direct URL/ID access."),
        ("Super Admin Portal Isolation", "Super Admin JWT cannot access enterprise ERP routes. Enterprise JWT cannot access Super Admin routes."),
        ("Reseller Tenant Isolation", "Reseller sees only their own tenants. Other resellers' tenants are completely hidden."),
        ("Reseller Max Tenant Limit", "When reseller reaches plan's max tenant count, creating another tenant is blocked with clear error."),
        ("Discount Validation", "Line item discount: 0% minimum, 100% maximum. Values outside this range rejected."),
        ("Quantity Validation", "All quantity fields must be positive integers. Zero and negative values rejected."),
        ("Approved Quotation Edit Lock", "Approved quotation cannot be directly edited. Must initiate Revise action for new version."),
    ],
    [2.0, 4.5]
)

# ── Section 6: Known Issues ───────────────────────────────────────────────────
add_h2(doc, "6. Known Issues / Current Bugs")
make_table(doc,
    ["#", "Module", "Issue Description", "Severity", "Status"],
    [
        ("BUG-001", "Quotation", "[PLACEHOLDER — e.g., Version increments on rejection instead of staying same]", "High", "Open"),
        ("BUG-002", "Purchase Order", "[PLACEHOLDER — e.g., Employee can view POs not assigned to them]", "High", "Open"),
        ("BUG-003", "Invoice", "[PLACEHOLDER — e.g., Invoice amount not auto-populated from PO total]", "Medium", "Open"),
        ("BUG-004", "RBAC", "[PLACEHOLDER — e.g., Manager can see data from other manager's team]", "High", "Open"),
        ("BUG-005", "Super Admin", "[PLACEHOLDER — fill in during testing]", "—", "Open"),
        ("BUG-006", "Reseller", "[PLACEHOLDER — fill in during testing]", "—", "Open"),
        ("BUG-007", "UI", "[PLACEHOLDER — e.g., Double-click on submit creates duplicate records]", "Medium", "Open"),
        ("BUG-008", "General", "[PLACEHOLDER — fill in during testing]", "—", "Open"),
    ],
    [0.7, 1.2, 3.3, 0.8, 0.7]
)
add_note(doc, "QA Analyst: Replace placeholders with actual defects found during testing. Log each defect with: module, steps to reproduce, expected vs actual result, severity, screenshot.")

# ── Section 7: Test Environment ───────────────────────────────────────────────
add_h2(doc, "7. Test Environment Setup")

add_h3(doc, "Application URLs")
make_table(doc,
    ["Portal", "URL"],
    [
        ("Main Application", "http://64.235.43.187:2262"),
        ("API Base URL", "http://64.235.43.187:2261"),
        ("Super Admin Login", "http://64.235.43.187:2262/superadmin/login"),
        ("Reseller Login", "http://64.235.43.187:2262/reseller/login"),
        ("Enterprise Login", "http://64.235.43.187:2262/login"),
    ],
    [2.0, 4.5]
)

add_h3(doc, "Test Credentials — All Roles")
make_table(doc,
    ["Role", "Portal URL", "Email", "Password"],
    [
        ("Super Admin", "/superadmin/login", "[SUPER_ADMIN_EMAIL]", "[SUPER_ADMIN_PASSWORD]"),
        ("Reseller Admin", "/reseller/login", "[RESELLER_EMAIL]", "[RESELLER_PASSWORD]"),
        ("Enterprise Admin", "/login", "[ADMIN_EMAIL]", "[ADMIN_PASSWORD]"),
        ("Manager A", "/login", "[MANAGER_A_EMAIL]", "[PASSWORD]"),
        ("Manager B", "/login", "[MANAGER_B_EMAIL]", "[PASSWORD]"),
        ("Employee X (under Mgr A)", "/login", "[EMP_X_EMAIL]", "[PASSWORD]"),
        ("Employee Y (under Mgr A)", "/login", "[EMP_Y_EMAIL]", "[PASSWORD]"),
        ("Employee Z (under Mgr B)", "/login", "[EMP_Z_EMAIL]", "[PASSWORD]"),
    ],
    [1.5, 1.8, 2.0, 1.2]
)
add_note(doc, "IMPORTANT: Replace all [PLACEHOLDER] values with actual credentials before testing. Create all accounts before the first test session.")

add_h3(doc, "Other Configuration")
make_table(doc,
    ["Parameter", "Value"],
    [
        ("Recommended Browser", "Google Chrome (latest)"),
        ("Secondary Browsers", "Firefox, Microsoft Edge"),
        ("Postman Collection", "[POSTMAN_COLLECTION_LINK]"),
        ("API Auth — Enterprise", "POST http://64.235.43.187:2261/auth/login  { email, password }"),
        ("API Auth — Super Admin", "POST http://64.235.43.187:2261/super-admin/auth/login  { email, password }"),
        ("API Auth — Reseller", "POST http://64.235.43.187:2261/resellers/auth/login  { email, password }"),
    ],
    [2.0, 4.5]
)

add_h3(doc, "Pre-Test Checklist")
for item in [
    "Super Admin account created and accessible at /superadmin/login",
    "Reseller account created under Super Admin and accessible at /reseller/login",
    "At least 1 tenant enterprise created by Reseller with active subscription",
    "All 5 enterprise user accounts created (Admin, Manager A, Manager B, Employee X, Y, Z)",
    "Manager A has Employee X and Employee Y assigned as direct reports",
    "Manager B has Employee Z assigned as direct report",
    "At least 3 vendors created in the enterprise",
    "At least 3 customers created in the enterprise",
    "At least 5 product/item records created",
    "Postman installed with 5 environments configured (one per role type)",
    "Browser dev tools available for console error monitoring",
]:
    add_bullet(doc, item)

# ── Section 8: Test Data ──────────────────────────────────────────────────────
add_h2(doc, "8. Test Data Requirements")
make_table(doc,
    ["Data Category", "Minimum Required", "Notes"],
    [
        ("Super Admin Account", "1 Super Admin", "Full platform access — created at server level"),
        ("Reseller Accounts", "1 Reseller", "Created by Super Admin"),
        ("Enterprise Tenants", "1 active Enterprise", "Created by Reseller with active subscription"),
        ("Enterprise User Accounts", "1 Admin, 2 Managers, 3 Employees", "Employees split across 2 managers (2+1)"),
        ("Vendors", "3 vendors", "With name, contact, and address"),
        ("Customers", "3 customers", "With name and billing address"),
        ("Products / Line Items", "5–10 products", "With unit prices defined"),
        ("RFQs", "5 RFQs", "Mix of Draft, Sent, Received, Converted states"),
        ("Quotations", "8 quotations", "Mix of v1, v2, Draft, Submitted, Approved, Rejected"),
        ("Purchase Orders", "4 POs", "Assigned: 2 to Emp X, 1 to Emp Y, 1 to Emp Z"),
        ("Invoices", "3 invoices", "Mix of Unpaid, Partially Paid, Paid statuses"),
        ("Sales Orders", "3 sales orders", "Assigned to different employees"),
    ],
    [1.8, 1.9, 2.8]
)

# ── Section 9: API Testing ────────────────────────────────────────────────────
add_h2(doc, "9. API Testing (Postman)")

add_h3(doc, "Setup — Three Portal Environments")
for b in [
    "Enterprise Base URL: http://64.235.43.187:2261",
    "Super Admin Auth: POST /super-admin/auth/login  |  Token type: super_admin",
    "Reseller Auth: POST /resellers/auth/login  |  Token type: reseller",
    "Enterprise Auth: POST /auth/login  |  Token type: enterprise user",
    "All endpoints require: Authorization: Bearer <token> header",
    "Create 5 Postman environments: Super Admin, Reseller, Enterprise Admin, Manager, Employee",
]:
    add_bullet(doc, b)

add_h3(doc, "Key Endpoints to Test")
make_table(doc,
    ["Method", "Endpoint", "What to Verify"],
    [
        ("POST", "/super-admin/auth/login", "Returns super admin JWT. Type = super_admin in token payload"),
        ("POST", "/resellers/auth/login", "Returns reseller JWT. Type = reseller in token payload"),
        ("POST", "/auth/login", "Returns enterprise JWT. Role returned in response"),
        ("GET", "/super-admin/enterprises", "Only super admin JWT succeeds. Enterprise JWT → 403"),
        ("GET", "/resellers/me/tenants", "Only reseller JWT succeeds. Super admin JWT → 403"),
        ("GET", "/quotations", "Admin = all, Manager = team, Employee = own only"),
        ("POST", "/quotations", "Creates at v1. Returns quotation ID"),
        ("PATCH", "/quotations/:id/submit", "Status → Submitted"),
        ("PATCH", "/quotations/:id/approve", "Status → Approved. Version unchanged"),
        ("PATCH", "/quotations/:id/reject", "Status → Rejected. Version must NOT change"),
        ("PATCH", "/quotations/:id/revise", "Status → Draft. Version must increment"),
        ("POST", "/purchase-orders", "Only Manager/Admin token should succeed (403 for Employee)"),
        ("GET", "/purchase-orders", "Employee sees only assigned. Manager sees team"),
        ("POST", "/invoices", "Blocked if PO not in Received state"),
        ("PATCH", "/invoices/:id/payment", "Records payment, updates status correctly"),
    ],
    [0.7, 1.8, 4.0]
)

# ── Section 10: QA Scope ──────────────────────────────────────────────────────
add_h2(doc, "10. QA Scope")
add_h3(doc, "In Scope")
for item in [
    "Functional testing of all 3 portals (Super Admin, Reseller, Enterprise)",
    "Super Admin: enterprise management, reseller management, subscriptions, support tickets",
    "Reseller Admin: tenant management, subscription plans, wallet, commissions",
    "Enterprise: all ERP modules (RFQ, Quotations, POs, Invoices, Sales, Employee Management)",
    "Role-based access control — all 5 roles tested independently",
    "Portal isolation — JWT cross-portal access must fail with 403",
    "Manager hierarchy visibility — cross-team data isolation verified",
    "Quotation version control logic (rejection vs revision behavior)",
    "End-to-end workflow: RFQ → Quotation → PO → Invoice",
    "End-to-end workflow: Enquiry → Sales Order → Invoice",
    "Negative testing (invalid inputs, unauthorized access, duplicate creation)",
    "Boundary testing (0%, 100%, 101% discount; zero/negative quantities; max tenant limits)",
    "UI behavior (form validation, button states, double-click prevention, navigation) — all portals",
    "API endpoint testing via Postman (auth for all 3 portals, scoping, status transitions)",
    "Cross-browser testing (Chrome primary, Firefox and Edge secondary)",
]:
    add_bullet(doc, "YES  " + item)

add_h3(doc, "High Priority Focus Areas")
for item in [
    "Quotation version control — rejection must NOT bump version",
    "PO visibility scoping by role and assignment",
    "RBAC enforcement — both UI and API layer — all 3 portals",
    "Manager hierarchy isolation — no cross-team data leakage",
    "Invoice creation gates — only from Received PO",
    "Duplicate record prevention",
    "Super Admin portal isolation — super admin JWT must not work on enterprise routes",
    "Reseller tenant isolation — reseller must not see other resellers' tenants",
    "Reseller max tenant limit enforcement",
]:
    add_bullet(doc, "STAR  " + item)

# ── Section 11: Out of Scope ──────────────────────────────────────────────────
add_h2(doc, "11. Out of Scope")
for item in [
    "Performance / load testing",
    "Stress testing or concurrent user simulation",
    "Mobile native app (web-responsive only in scope)",
    "Email notification content/delivery testing",
    "Third-party integrations",
    "Database backup and recovery testing",
    "Infrastructure / server configuration testing",
    "Accessibility (WCAG) compliance testing",
]:
    add_bullet(doc, "NO  " + item)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  PART 2 — TEST CASES (185 Total)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, "PART 2 — TEST CASES  (185 Total)")
add_note(doc, "Columns: TC ID | Module | Test Scenario | Test Steps | Expected Result | Actual Result | Status | Priority")

SECTIONS = {
    "A": "Section A — Authentication — All Roles (TC-001 to TC-018)",
    "B": "Section B — Super Admin Module (TC-019 to TC-040)",
    "C": "Section C — Reseller Admin Module (TC-041 to TC-058)",
    "D": "Section D — Role-Based Access Control / RBAC (TC-059 to TC-067)",
    "E": "Section E — Manager Hierarchy & Visibility (TC-068 to TC-073)",
    "F": "Section F — RFQ Module (TC-074 to TC-086)",
    "G": "Section G — Quotation Module (TC-087 to TC-110)",
    "H": "Section H — Purchase Order Module (TC-111 to TC-125)",
    "I": "Section I — Invoice Module (TC-126 to TC-139)",
    "J": "Section J — Sales Module (TC-140 to TC-147)",
    "K": "Section K — Employee & Role Management (TC-148 to TC-157)",
    "L": "Section L — UI, Navigation & Forms (TC-158 to TC-169)",
    "M": "Section M — API Testing (TC-170 to TC-180)",
    "N": "Section N — Negative & Edge Cases (TC-181 to TC-185)",
}

PRIORITY_RGB = {
    "High":   RGBColor(0xFE, 0xE2, 0xE2),
    "Medium": RGBColor(0xFE, 0xF9, 0xC3),
    "Low":    RGBColor(0xDC, 0xFC, 0xE7),
}

test_cases = [
    # SECTION A — Authentication All Roles
    ("A","TC-001","Auth","Super Admin login with valid credentials","1. Go to /superadmin/login\n2. Enter super admin email & password\n3. Click Login","Super Admin dashboard loads with full platform menu (Enterprises, Resellers, Subscriptions, Support, etc.)","","","High"),
    ("A","TC-002","Auth","Reseller Admin login with valid credentials","1. Go to /reseller/login\n2. Enter reseller email & password\n3. Click Login","Reseller dashboard loads with tenant management menu (Tenants, Plans, Wallet, Commissions, etc.)","","","High"),
    ("A","TC-003","Auth","Enterprise Admin login with valid credentials","1. Go to /login\n2. Enter admin email & password\n3. Click Login","Admin dashboard loads with full enterprise menu","","","High"),
    ("A","TC-004","Auth","Manager login with valid credentials","1. Go to /login\n2. Enter manager credentials\n3. Click Login","Manager dashboard loads; team-scoped menu visible","","","High"),
    ("A","TC-005","Auth","Employee login with valid credentials","1. Go to /login\n2. Enter employee credentials\n3. Click Login","Employee dashboard loads; restricted menu (no PO create, no approvals)","","","High"),
    ("A","TC-006","Auth","Super Admin login with wrong password","1. Go to /superadmin/login\n2. Enter wrong password\n3. Click Login","Error: Invalid credentials. No redirect. Token not issued","","","High"),
    ("A","TC-007","Auth","Reseller login with wrong password","1. Go to /reseller/login\n2. Enter wrong password\n3. Click Login","Error: Invalid credentials. No access granted","","","High"),
    ("A","TC-008","Auth","Super Admin token cannot access Enterprise routes","1. Login as Super Admin\n2. Call GET /quotations (enterprise API) with super admin JWT","401 or 403 — Super admin JWT is not valid on enterprise routes","","","High"),
    ("A","TC-009","Auth","Reseller token cannot access Super Admin routes","1. Login as Reseller\n2. Call GET /super-admin/enterprises with reseller JWT","403 Forbidden — ResellerGuard does not pass SuperAdminGuard","","","High"),
    ("A","TC-010","Auth","Employee token cannot access Reseller routes","1. Login as Employee\n2. Call GET /resellers/me/tenants with employee JWT","403 Forbidden — wrong token type","","","High"),
    ("A","TC-011","Auth","Login with blank fields — all portals","1. Open any login page\n2. Leave email and password blank\n3. Click Login","Frontend validation errors shown on both fields. No API call made","","","Medium"),
    ("A","TC-012","Auth","Session persistence after refresh — Super Admin","1. Login as Super Admin\n2. Refresh browser","Super Admin remains logged in. Dashboard reloads correctly","","","High"),
    ("A","TC-013","Auth","Session persistence after refresh — Reseller","1. Login as Reseller\n2. Refresh browser","Reseller remains logged in. Dashboard reloads correctly","","","High"),
    ("A","TC-014","Auth","Logout clears session — Super Admin","1. Login as Super Admin\n2. Click Logout","Session cleared. Redirected to /superadmin/login. Back button does not return","","","High"),
    ("A","TC-015","Auth","Logout clears session — Reseller","1. Login as Reseller\n2. Click Logout","Session cleared. Redirected to /reseller/login","","","High"),
    ("A","TC-016","Auth","Locked account cannot login","1. Super Admin locks a reseller account\n2. That reseller tries to login","Login rejected: Account is locked or inactive","","","High"),
    ("A","TC-017","Auth","SQL injection in Super Admin login","1. Go to /superadmin/login\n2. Enter ' OR 1=1 -- in email field","Login rejected. No data exposed. Standard error shown","","","High"),
    ("A","TC-018","Auth","Expired reseller subscription restricts features","1. Reseller subscription is expired\n2. Login as that Reseller","Login succeeds but shows subscription expired warning. Tenant creation blocked","","","Medium"),
    # SECTION B — Super Admin
    ("B","TC-019","Super Admin","Super Admin dashboard loads platform KPIs","1. Login as Super Admin\n2. Open dashboard","Dashboard shows: Total Enterprises, Active Enterprises, Total Resellers, Platform Revenue, Monthly Charts","","","High"),
    ("B","TC-020","Super Admin","Super Admin views all enterprises","1. Login as Super Admin\n2. Navigate to Enterprises","All enterprises listed with name, status, subscription expiry, creation date","","","High"),
    ("B","TC-021","Super Admin","Super Admin creates a new enterprise","1. Login as Super Admin\n2. Navigate to Enterprises\n3. Click Create\n4. Fill details and assign to Reseller\n5. Save","New enterprise created. Appears in enterprises list. Assigned reseller can see it","","","High"),
    ("B","TC-022","Super Admin","Super Admin views enterprise financial details","1. Login as Super Admin\n2. Open an enterprise\n3. Change period to 90 days","Revenue, costs, profit margins, invoice/PO/SO counts all shown for 90-day period","","","High"),
    ("B","TC-023","Super Admin","Super Admin approves an enterprise","1. Login as Super Admin\n2. Open a pending enterprise\n3. Click Approve","Enterprise status changes to Active. Enterprise admin can now login","","","High"),
    ("B","TC-024","Super Admin","Super Admin locks an enterprise","1. Login as Super Admin\n2. Open an active enterprise\n3. Click Lock","Enterprise status = Locked. Enterprise admin login is blocked","","","High"),
    ("B","TC-025","Super Admin","Super Admin unlocks an enterprise","1. Login as Super Admin\n2. Open a locked enterprise\n3. Click Unlock","Enterprise status = Active. Enterprise admin login is restored","","","High"),
    ("B","TC-026","Super Admin","Super Admin reassigns enterprise to different reseller","1. Login as Super Admin\n2. Open an enterprise\n3. Change reseller\n4. Save","Enterprise now appears under new reseller's tenant list. Previous reseller no longer sees it","","","High"),
    ("B","TC-027","Super Admin","Super Admin manages resellers list","1. Login as Super Admin\n2. Navigate to Resellers","All resellers listed with name, status, plan, wallet balance, tenant count","","","High"),
    ("B","TC-028","Super Admin","Super Admin creates a new reseller","1. Login as Super Admin\n2. Navigate to Resellers\n3. Click Create\n4. Fill reseller details\n5. Save","New reseller account created. Reseller can login at /reseller/login","","","High"),
    ("B","TC-029","Super Admin","Super Admin locks a reseller account","1. Login as Super Admin\n2. Open a reseller\n3. Click Lock","Reseller status = Locked. Reseller cannot login. Their tenants are unaffected","","","High"),
    ("B","TC-030","Super Admin","Super Admin credits reseller wallet","1. Login as Super Admin\n2. Open Resellers > Wallets\n3. Select reseller\n4. Credit 5000 to wallet","Reseller wallet balance increases by 5000. Transaction record created","","","High"),
    ("B","TC-031","Super Admin","Super Admin creates a subscription plan","1. Login as Super Admin\n2. Navigate to Subscriptions\n3. Create new plan with price, features, limits\n4. Save","Plan created and available for resellers to assign to tenants","","","High"),
    ("B","TC-032","Super Admin","Super Admin manages reseller plans","1. Login as Super Admin\n2. Navigate to Resellers > Plans\n3. Create plan with commission % and max tenants","Plan created. Resellers can subscribe to this plan","","","High"),
    ("B","TC-033","Super Admin","Super Admin views all support tickets","1. Login as Super Admin\n2. Navigate to Support","All support tickets from all enterprises and resellers listed","","","High"),
    ("B","TC-034","Super Admin","Super Admin replies to a support ticket","1. Login as Super Admin\n2. Open a ticket\n3. Type reply\n4. Submit","Reply saved. Ticket status updated. Requester notified","","","High"),
    ("B","TC-035","Super Admin","Super Admin manages services catalog","1. Login as Super Admin\n2. Navigate to Services\n3. Add/edit a service","Service catalog updated. Reflects across all enterprise plans","","","Medium"),
    ("B","TC-036","Super Admin","Super Admin creates a coupon code","1. Login as Super Admin\n2. Navigate to Coupons\n3. Create coupon with discount % and expiry\n4. Save","Coupon created. Can be applied during enterprise subscription","","","Medium"),
    ("B","TC-037","Super Admin","Super Admin views platform-level accounts","1. Login as Super Admin\n2. Navigate to Accounts\n3. Change period filter to 365 days","Platform-level revenue summary shown for selected period","","","Medium"),
    ("B","TC-038","Super Admin","Super Admin views all employees across platform","1. Login as Super Admin\n2. Navigate to Employees","All employees across all enterprises listed with name, role, and enterprise","","","Medium"),
    ("B","TC-039","Super Admin","Super Admin updates enterprise subscription expiry","1. Login as Super Admin\n2. Open an enterprise\n3. Update subscription expiry date","Expiry date saved. Enterprise subscription extended","","","High"),
    ("B","TC-040","Super Admin","Super Admin cannot access Enterprise ERP modules","1. Login as Super Admin\n2. Try to navigate to /quotations or /purchase-orders","Access denied. Super Admin portal is completely separate from ERP modules","","","High"),
    # SECTION C — Reseller Admin
    ("C","TC-041","Reseller Admin","Reseller dashboard loads KPIs","1. Login as Reseller\n2. Open dashboard","Shows: Total Tenants, Active Subscriptions, Expired, Revenue, Commission, Wallet Balance","","","High"),
    ("C","TC-042","Reseller Admin","Reseller views all their tenants","1. Login as Reseller\n2. Navigate to Tenants","Only enterprises assigned to this reseller are listed","","","High"),
    ("C","TC-043","Reseller Admin","Reseller creates a new tenant","1. Login as Reseller\n2. Navigate to Tenants\n3. Click Create\n4. Fill enterprise details\n5. Assign a plan\n6. Save","New tenant created under this reseller. Tenant admin can now login","","","High"),
    ("C","TC-044","Reseller Admin","Reseller cannot see other resellers tenants","1. Login as Reseller A\n2. Open Tenants","Only Reseller A's tenants visible. Reseller B's tenants NOT shown","","","High"),
    ("C","TC-045","Reseller Admin","Reseller assigns subscription plan to tenant","1. Login as Reseller\n2. Open a tenant\n3. Assign a subscription plan\n4. Save","Plan assigned to tenant. Subscription activated with expiry date","","","High"),
    ("C","TC-046","Reseller Admin","Reseller renews expiring tenant subscription","1. Login as Reseller\n2. Open Subscriptions list\n3. Find expiring tenant\n4. Click Renew","Subscription extended. New expiry date set. Wallet debited or invoice created","","","High"),
    ("C","TC-047","Reseller Admin","Reseller views wallet balance and transactions","1. Login as Reseller\n2. Navigate to Wallet","Wallet balance shown. Transaction history (credits, debits, commissions) listed","","","High"),
    ("C","TC-048","Reseller Admin","Reseller views commission earnings","1. Login as Reseller\n2. Navigate to Commissions","Commission earnings listed by tenant/plan/date. Total commission shown","","","High"),
    ("C","TC-049","Reseller Admin","Reseller sets custom pricing for a plan","1. Login as Reseller\n2. Navigate to Plans\n3. Select a plan\n4. Set custom price\n5. Save","Custom price saved. Tenants of this reseller are charged the custom price","","","High"),
    ("C","TC-050","Reseller Admin","Reseller views usage per tenant","1. Login as Reseller\n2. Navigate to Usage","Employee count, active vs total employees shown per tenant","","","Medium"),
    ("C","TC-051","Reseller Admin","Reseller views subscription status of all tenants","1. Login as Reseller\n2. Navigate to Subscriptions","All tenant subscriptions listed with status (Active, Expired, Expiring Soon)","","","High"),
    ("C","TC-052","Reseller Admin","Reseller views their own subscription plan","1. Login as Reseller\n2. Navigate to My Subscription","Own reseller plan details shown: plan name, max tenants, commission %, expiry","","","Medium"),
    ("C","TC-053","Reseller Admin","Reseller views billing/invoice history","1. Login as Reseller\n2. Navigate to Billing","All billing invoices listed with amounts and payment status","","","Medium"),
    ("C","TC-054","Reseller Admin","Reseller views reports","1. Login as Reseller\n2. Navigate to Reports","Sales/revenue/commission/tenant activity reports shown for own portfolio","","","Medium"),
    ("C","TC-055","Reseller Admin","Reseller updates own profile","1. Login as Reseller\n2. Navigate to Profile\n3. Update name and company\n4. Save","Profile updated. Changes saved and reflected in header/account area","","","Low"),
    ("C","TC-056","Reseller Admin","Reseller cannot access Super Admin portal","1. Login as Reseller\n2. Navigate to /superadmin/dashboard","Access denied or redirected to /reseller/login","","","High"),
    ("C","TC-057","Reseller Admin","Reseller cannot access Enterprise ERP modules","1. Login as Reseller\n2. Navigate to /quotations or /purchase-orders","Access denied. Reseller portal is completely separate from ERP modules","","","High"),
    ("C","TC-058","Reseller Admin","Reseller max tenant limit enforced","1. Reseller plan has max 5 tenants\n2. Reseller already has 5 tenants\n3. Try to create 6th tenant","Error: Maximum tenant limit reached for your plan. Upgrade plan to add more","","","High"),
    # SECTION D — RBAC
    ("D","TC-059","RBAC","Employee cannot access PO creation","1. Login as Employee\n2. Navigate to /purchase-orders/new","Access denied or button not visible. No PO form rendered","","","High"),
    ("D","TC-060","RBAC","Employee cannot generate invoices","1. Login as Employee\n2. Open a Received PO\n3. Check for Generate Invoice action","Generate Invoice button not visible to Employee","","","High"),
    ("D","TC-061","RBAC","Employee cannot access User Management","1. Login as Employee\n2. Navigate to /settings/users via URL","Redirect to dashboard or 403 error. No user list shown","","","High"),
    ("D","TC-062","RBAC","Manager cannot access User Management","1. Login as Manager\n2. Navigate to /settings/users","Access denied. Create/Edit user actions unavailable","","","High"),
    ("D","TC-063","RBAC","Manager can approve quotations","1. Login as Manager\n2. Open Submitted quotation\n3. Click Approve","Approval action succeeds. Status changes to Approved","","","High"),
    ("D","TC-064","RBAC","Employee cannot approve quotations","1. Login as Employee\n2. Open a Submitted quotation","Approve button not visible. Cannot trigger approval action","","","High"),
    ("D","TC-065","RBAC","Admin can access all ERP modules","1. Login as Enterprise Admin\n2. Navigate through every menu item","All ERP modules accessible without restriction","","","High"),
    ("D","TC-066","RBAC","Direct URL access by unauthorized role","1. Login as Employee\n2. Manually type /purchase-orders/new in URL bar","Redirected away or 403 error shown. No form rendered","","","High"),
    ("D","TC-067","RBAC","API call with Employee token to Manager-only endpoint","1. Login as Employee via Postman\n2. Call POST /purchase-orders with employee JWT","403 Forbidden response returned","","","High"),
    # SECTION E — Manager Hierarchy
    ("E","TC-068","Hierarchy","Manager A sees Employee X and Y records","1. Employee X and Y (both under Manager A) create quotations\n2. Login as Manager A","Records of both Employee X and Employee Y visible to Manager A","","","High"),
    ("E","TC-069","Hierarchy","Manager A cannot see Employee Z records (under B)","1. Employee Z (under Manager B) creates a quotation\n2. Login as Manager A","Employee Z quotation NOT in Manager A list","","","High"),
    ("E","TC-070","Hierarchy","Manager B cannot see Manager A team records","1. Employee X (under Manager A) creates a record\n2. Login as Manager B","Employee X record not visible to Manager B","","","High"),
    ("E","TC-071","Hierarchy","Admin sees all records from all users","1. Employees under Manager A and B each create records\n2. Login as Admin","All records from all users visible","","","High"),
    ("E","TC-072","Hierarchy","Employee reassigned — visibility updates immediately","1. Employee X under Manager A\n2. Admin reassigns to Manager B\n3. Login as Manager B","Employee X records visible to Manager B. No longer visible to Manager A","","","High"),
    ("E","TC-073","Hierarchy","Manager own records visible to themselves","1. Manager A creates a quotation directly\n2. Login as Manager A","Manager A own quotation visible in their list","","","Medium"),
    # SECTION F — RFQ
    ("F","TC-074","RFQ","Create RFQ with valid data","1. Login\n2. Navigate to RFQ\n3. Fill vendor, items, quantities\n4. Submit","RFQ created with status Draft/Sent and unique reference number","","","High"),
    ("F","TC-075","RFQ","Create RFQ with missing vendor","1. Fill RFQ items\n2. Leave vendor blank\n3. Submit","Validation error: Vendor is required","","","High"),
    ("F","TC-076","RFQ","Create RFQ with zero quantity","1. Create RFQ\n2. Set item qty = 0\n3. Submit","Validation error: Quantity must be greater than 0","","","Medium"),
    ("F","TC-077","RFQ","Create RFQ with negative quantity","1. Create RFQ\n2. Enter qty = -5\n3. Submit","Validation error: negative values not allowed","","","Medium"),
    ("F","TC-078","RFQ","Create RFQ with no line items","1. Fill vendor\n2. Add no items\n3. Submit","Validation error: At least one line item is required","","","High"),
    ("F","TC-079","RFQ","Convert RFQ to Quotation","1. Open a submitted/received RFQ\n2. Click Convert to Quotation","Quotation created at v1 with all items pre-filled from RFQ","","","High"),
    ("F","TC-080","RFQ","Quotation from RFQ pre-fills all line items","1. Create RFQ with 3 items\n2. Convert to Quotation\n3. Open the quotation","All 3 items present in quotation with matching quantity and description","","","High"),
    ("F","TC-081","RFQ","Employee sees only their own RFQs","1. Employee X creates an RFQ\n2. Login as Employee Y","Employee Y cannot see Employee X RFQ","","","High"),
    ("F","TC-082","RFQ","Manager sees all team RFQs","1. Employee X (under Manager A) creates an RFQ\n2. Login as Manager A","RFQ appears in Manager A list","","","High"),
    ("F","TC-083","RFQ","Duplicate RFQ same vendor same reference","1. Create RFQ for Vendor A ref #001\n2. Create another with same vendor and reference","System warns or blocks exact duplicate","","","Medium"),
    ("F","TC-084","RFQ","RFQ list displays correct status badges","1. Create RFQs in different states\n2. Open RFQ list","Status badges correctly show Draft, Sent, Received per record","","","Low"),
    ("F","TC-085","RFQ","Delete a Draft RFQ","1. Create an RFQ\n2. Keep it in Draft\n3. Delete it","RFQ removed from list. No linked records affected","","","Medium"),
    ("F","TC-086","RFQ","Cannot delete a converted RFQ","1. Convert RFQ to Quotation\n2. Try to delete the original RFQ","Deletion blocked: RFQ has a linked Quotation","","","Medium"),
    # SECTION G — Quotation
    ("G","TC-087","Quotation","Create quotation with valid line items","1. Navigate to Create Quotation\n2. Add items with qty, unit price, discount\n3. Save","Quotation saved at v1, status = Draft, total calculated correctly","","","High"),
    ("G","TC-088","Quotation","Total calculation — single item with discount","1. Add Item: qty=2, price=1000, discount=10%","Total = (2x1000) - 10% = 1800","","","High"),
    ("G","TC-089","Quotation","Total calculation — multiple items","1. Item A: qty=1, price=500, 0% discount\n2. Item B: qty=2, price=300, 5% discount","Total = 500 + 570 = 1070","","","High"),
    ("G","TC-090","Quotation","Discount over 100%","1. Create quotation\n2. Enter discount = 101%\n3. Save","Validation error: Discount cannot exceed 100%","","","High"),
    ("G","TC-091","Quotation","Create quotation with no line items","1. Open Create Quotation\n2. Add no line items\n3. Submit","Validation error: At least one line item is required","","","High"),
    ("G","TC-092","Quotation","Submit quotation for approval","1. Open a Draft quotation\n2. Click Submit","Status = Submitted. Approve/Reject options now available to approvers","","","High"),
    ("G","TC-093","Quotation","Approve a submitted quotation","1. Login as Manager or Admin\n2. Open Submitted quotation\n3. Click Approve","Status = Approved. Version number unchanged. Edit fields locked","","","High"),
    ("G","TC-094","Quotation","Reject a submitted quotation","1. Login as Manager or Admin\n2. Open Submitted quotation\n3. Click Reject\n4. Enter rejection reason","Status = Rejected. Version number stays the same (v1 stays v1)","","","High"),
    ("G","TC-095","Quotation","CRITICAL: Version does NOT increment on rejection","1. Create Quotation (version = 1)\n2. Submit\n3. Manager rejects\n4. Check version field","Version field = 1 unchanged. Rejection must NEVER auto-increment version","","","High"),
    ("G","TC-096","Quotation","Revise a rejected quotation","1. Open Rejected quotation\n2. Click Revise\n3. Modify at least one line item\n4. Resubmit","New version created: v1 to v2. Original v1 preserved in version history","","","High"),
    ("G","TC-097","Quotation","CRITICAL: Version increments on user revision","1. Reject a v1 quotation\n2. User clicks Revise\n3. Check version after save","Version = 2. Re-submission shows v2 to approver","","","High"),
    ("G","TC-098","Quotation","Multiple revision cycles — correct versioning","1. v1 Reject Revise v2 Reject Revise v3","Versions: v1, v2, v3 in correct sequence. All listed in history","","","High"),
    ("G","TC-099","Quotation","Approved quotation fields are read-only","1. Open an Approved quotation\n2. Try to click and edit any field","All fields read-only. Only Revise action button available","","","High"),
    ("G","TC-100","Quotation","Revise approved quotation creates new version","1. Open Approved quotation\n2. Click Revise\n3. Edit and resubmit","New version (e.g. v2) created in Draft. Previous approved version preserved","","","High"),
    ("G","TC-101","Quotation","Employee cannot approve own submitted quotation","1. Login as Employee\n2. Create and submit a quotation\n3. Check for Approve button","Approve button NOT visible to the Employee who submitted it","","","High"),
    ("G","TC-102","Quotation","Quotation linked to correct customer","1. Create quotation for Customer A\n2. Open quotation detail","Customer A shown in customer field. No data leakage from other customers","","","High"),
    ("G","TC-103","Quotation","Search quotation by customer name","1. Open Quotation list\n2. Search by customer name","Only quotations linked to that customer are shown","","","Medium"),
    ("G","TC-104","Quotation","Filter quotations by status","1. Open Quotation list\n2. Apply filter Approved","Only Approved quotations displayed in list","","","Medium"),
    ("G","TC-105","Quotation","Cannot re-submit already Submitted quotation","1. Submit a quotation\n2. Click Submit again","Submit button disabled or already-submitted message shown","","","Medium"),
    ("G","TC-106","Quotation","Quotation PDF/print generation","1. Open an Approved quotation\n2. Click Download or Print","PDF generated with correct line items, totals, customer details, and version","","","Medium"),
    ("G","TC-107","Quotation","Duplicate quotation reference prevention","1. Create Quotation ref #Q-001 for Customer A\n2. Attempt to create another with same ref/customer","System warns or prevents exact duplicate quotation creation","","","Medium"),
    ("G","TC-108","Quotation","Version history shows all revisions with metadata","1. Create quotation with 3 revision cycles\n2. Open Version History tab","v1, v2, v3 listed with timestamp, status, and who approved/rejected","","","Medium"),
    ("G","TC-109","Quotation","Discount exactly 100%","1. Create quotation\n2. Enter discount = 100% on a line item","Accepted — line item cost = 0. Total reflects correctly","","","Medium"),
    ("G","TC-110","Quotation","Discount = 0% no discount","1. Leave discount blank or enter 0","No discount applied. Full price used in total","","","Low"),
    # SECTION H — Purchase Order
    ("H","TC-111","Purchase Order","Create PO from approved quotation","1. Open Approved quotation\n2. Click Create PO\n3. Assign to Employee\n4. Save","PO created with correct amounts and assigned to selected employee","","","High"),
    ("H","TC-112","Purchase Order","PO amount matches approved quotation total","1. Approved quotation total = 25000\n2. Create PO from it","PO total = 25000. No discrepancy between quotation and PO","","","High"),
    ("H","TC-113","Purchase Order","Cannot create PO from Draft quotation","1. Open a Draft quotation\n2. Check for Create PO option","Create PO button not visible or disabled","","","High"),
    ("H","TC-114","Purchase Order","Cannot create PO from Rejected quotation","1. Open a Rejected quotation\n2. Check for Create PO option","Create PO button not available for rejected quotations","","","High"),
    ("H","TC-115","Purchase Order","Cannot create PO from Submitted quotation","1. Open a Submitted quotation\n2. Check Create PO","Create PO disabled — must wait for approval","","","High"),
    ("H","TC-116","Purchase Order","Duplicate PO prevention for same quotation","1. Create PO from Approved quotation\n2. Try to create another PO from same quotation","System blocks or warns: A PO already exists for this quotation","","","High"),
    ("H","TC-117","Purchase Order","Employee sees only their assigned POs","1. Create PO assigned to Employee X\n2. Login as Employee X\n3. Open PO list","PO assigned to Employee X is visible","","","High"),
    ("H","TC-118","Purchase Order","Employee cannot see PO assigned to another employee","1. Create PO assigned to Employee Y\n2. Login as Employee X\n3. Open PO list","Employee Y PO is NOT in Employee X list","","","High"),
    ("H","TC-119","Purchase Order","Manager sees all team POs","1. Create POs for Employee X and Y (both under Manager A)\n2. Login as Manager A","Both POs from Employee X and Y visible to Manager A","","","High"),
    ("H","TC-120","Purchase Order","Update PO status to Received","1. Open an active PO\n2. Update status to Received","Status updates to Received. Invoice generation now enabled","","","High"),
    ("H","TC-121","Purchase Order","PO cannot be deleted if linked invoice exists","1. Generate an invoice from PO\n2. Try to delete the PO","Deletion blocked: Cannot delete a PO with an associated invoice","","","High"),
    ("H","TC-122","Purchase Order","Search PO by vendor name","1. Open PO list\n2. Search Vendor A","Only POs linked to Vendor A shown","","","Medium"),
    ("H","TC-123","Purchase Order","PO reference number is unique","1. Create multiple POs","Each PO has a unique system-generated reference number","","","High"),
    ("H","TC-124","Purchase Order","Reassign PO updates visibility","1. PO assigned to Employee X\n2. Edit PO and reassign to Employee Y","PO now visible to Employee Y. No longer visible to Employee X","","","Medium"),
    ("H","TC-125","Purchase Order","PO list sorted by date newest first","1. Create 3 POs on different dates\n2. Open PO list","Most recently created PO appears first","","","Low"),
    # SECTION I — Invoice
    ("I","TC-126","Invoice","Generate invoice from Received PO","1. Open PO with status = Received\n2. Click Generate Invoice","Invoice created with amount pre-filled from PO. Status = Unpaid","","","High"),
    ("I","TC-127","Invoice","Invoice amount auto-populated from PO","1. PO total = 40000\n2. Generate invoice","Invoice amount = 40000. No manual entry required","","","High"),
    ("I","TC-128","Invoice","Invoice default status is Unpaid","1. Generate any new invoice","Invoice status = Unpaid on creation","","","High"),
    ("I","TC-129","Invoice","Record partial payment","1. Invoice total = 50000\n2. Record payment of 20000","Status = Partially Paid. Remaining balance = 30000 shown","","","High"),
    ("I","TC-130","Invoice","Record full payment","1. Invoice total = 50000\n2. Record payment of 50000","Status = Paid. Balance = 0. Invoice marked complete","","","High"),
    ("I","TC-131","Invoice","Payment amount exceeds invoice total","1. Invoice = 50000\n2. Attempt to record 60000 payment","Validation error: Payment cannot exceed invoice total","","","High"),
    ("I","TC-132","Invoice","Multiple partial payments until fully paid","1. Invoice = 50000\n2. Record 20000 then 20000 then 10000","Each payment recorded. Final payment triggers status = Paid","","","High"),
    ("I","TC-133","Invoice","Employee cannot generate invoice","1. Login as Employee\n2. Open a Received PO","Generate Invoice button not visible to Employee role","","","High"),
    ("I","TC-134","Invoice","Invoice linked to correct PO","1. Generate invoice from PO #P-001\n2. Open invoice detail","Invoice shows reference to PO #P-001 with a working link","","","Medium"),
    ("I","TC-135","Invoice","Filter invoice list by payment status","1. Open Invoices list\n2. Filter by Paid","Only Paid invoices shown in list","","","Medium"),
    ("I","TC-136","Invoice","Invoice PDF/print generation","1. Open an invoice\n2. Click Download or Print","PDF generated with correct amounts, payment status, and customer details","","","Medium"),
    ("I","TC-137","Invoice","Paid invoice is read-only","1. Open invoice with status = Paid\n2. Try to modify any field or add a payment","All fields and actions locked. No modification possible","","","High"),
    ("I","TC-138","Invoice","Duplicate invoice prevention for same PO","1. Generate invoice from PO #P-001\n2. Try to generate another invoice from same PO","Blocked: An invoice already exists for this Purchase Order","","","High"),
    ("I","TC-139","Invoice","Generate invoice from PO not yet Received","1. Open a PO with status = Ordered\n2. Attempt to generate invoice","Invoice generation blocked: PO must be in Received status","","","High"),
    # SECTION J — Sales
    ("J","TC-140","Sales","Create customer sales order with valid data","1. Navigate to Sales\n2. Create order for Customer A\n3. Add line items\n4. Save","Sales order saved with unique reference and correct status","","","High"),
    ("J","TC-141","Sales","Sales order linked to approved quotation","1. Approve a customer quotation\n2. Convert to Sales Order","Sales order references the originating quotation. Line items match","","","High"),
    ("J","TC-142","Sales","Sales order triggers job card creation","1. Create a Sales Order\n2. Click Create Job Card","Job card created with reference back to the sales order","","","High"),
    ("J","TC-143","Sales","Employee sees only assigned sales orders","1. Create 2 sales orders (1 to X, 1 to Y)\n2. Login as Employee X","Only Employee X sales order visible in their list","","","High"),
    ("J","TC-144","Sales","Manager sees all team sales orders","1. Sales orders assigned to Employee X and Y (both under Manager A)\n2. Login as Manager A","Both sales orders visible to Manager A","","","High"),
    ("J","TC-145","Sales","Update sales order status","1. Open a sales order\n2. Change status from In Progress to Dispatched","Status updates correctly. History log reflects the change with timestamp","","","Medium"),
    ("J","TC-146","Sales","Sales order total calculation","1. Add 2 line items with qty, price, and discount\n2. Check total","Total reflects correct arithmetic with all discounts applied","","","High"),
    ("J","TC-147","Sales","Cancel a sales order","1. Open an In Progress sales order\n2. Click Cancel","Status = Cancelled. Associated job card status also updated to Cancelled","","","Medium"),
    # SECTION K — Employee Management
    ("K","TC-148","Employee Mgmt","Admin creates new employee user","1. Login as Admin\n2. Navigate to Users\n3. Create user role = Employee\n4. Assign to Manager A","User created. Appears in Manager A team immediately","","","High"),
    ("K","TC-149","Employee Mgmt","Admin creates new manager user","1. Login as Admin\n2. Create user with role = Manager","Manager account created. Can see team records once employees assigned","","","High"),
    ("K","TC-150","Employee Mgmt","Admin assigns employee to a manager","1. Login as Admin\n2. Edit Employee X\n3. Set reporting manager = Manager A","Employee X appears under Manager A team","","","High"),
    ("K","TC-151","Employee Mgmt","Admin reassigns employee to new manager","1. Employee X under Manager A\n2. Admin reassigns to Manager B","Manager B now sees Employee X records. Manager A no longer sees them","","","High"),
    ("K","TC-152","Employee Mgmt","Manager cannot create user accounts","1. Login as Manager\n2. Navigate to User Management","Create/Edit user actions not available to Manager role","","","High"),
    ("K","TC-153","Employee Mgmt","Admin upgrades employee role to Manager","1. Login as Admin\n2. Edit Employee X\n3. Change role to Manager","Employee X gets Manager-level access on next login","","","Medium"),
    ("K","TC-154","Employee Mgmt","Admin deactivates a user account","1. Login as Admin\n2. Deactivate Employee X","Employee X cannot login. Existing records remain visible to Manager/Admin","","","Medium"),
    ("K","TC-155","Employee Mgmt","Deactivated user cannot login","1. Admin deactivates Employee X\n2. Employee X attempts to login","Login rejected: Account deactivated or Invalid credentials","","","High"),
    ("K","TC-156","Employee Mgmt","Admin edits user profile details","1. Login as Admin\n2. Edit name/email of a user\n3. Save","Changes saved. Updated information reflects across the system","","","Medium"),
    ("K","TC-157","Employee Mgmt","Duplicate email on user creation","1. Login as Admin\n2. Create user with email already registered","Validation error: Email already registered","","","High"),
    # SECTION L — UI
    ("L","TC-158","UI / Navigation","Sidebar navigation — all links work (all portals)","1. Login as each role type\n2. Click each sidebar menu item","Correct page loads for each role. No 404 errors encountered","","","High"),
    ("L","TC-159","UI / Navigation","Breadcrumb navigation accuracy","1. Navigate to a nested page\n2. Check breadcrumb","Breadcrumb shows correct path","","","Low"),
    ("L","TC-160","UI / Navigation","Back button preserves list filter state","1. Apply a filter on quotation list\n2. Open a record\n3. Press browser back","Returns to list with same filter still applied","","","Low"),
    ("L","TC-161","UI / Forms","Double-click submit prevention","1. Fill a quotation form completely\n2. Rapidly double-click Submit button","Only one record created. No duplicate quotation in the list","","","High"),
    ("L","TC-162","UI / Forms","Required field highlighting on submit","1. Open any create form\n2. Leave required fields blank\n3. Click Submit","All blank required fields highlighted in red with descriptive error text","","","High"),
    ("L","TC-163","UI / Forms","Form data persists on browser tab switch","1. Start filling a form\n2. Switch to another browser tab\n3. Return","Form data still intact. No data lost","","","Medium"),
    ("L","TC-164","UI / Forms","Long text input in description field","1. Enter 1000+ characters in any description field\n2. Save","Content accepted and saved. Displayed correctly on detail view","","","Low"),
    ("L","TC-165","UI / Forms","XSS injection in text fields","1. Enter script tag in a name or description field\n2. Save and reload","Input sanitized. Script does not execute. Stored as literal text","","","High"),
    ("L","TC-166","UI / Tables","Pagination on large dataset","1. Login as Admin\n2. Open list with 50+ records","Data loads in pages. Pagination controls work correctly","","","Medium"),
    ("L","TC-167","UI / Tables","Column sorting works correctly","1. Open any list table\n2. Click a column header","List re-orders ascending then descending on repeated clicks","","","Medium"),
    ("L","TC-168","UI / Tables","Export to Excel or CSV","1. Open any list\n2. Click Export button","File downloaded with correct column headers and all visible data","","","Medium"),
    ("L","TC-169","UI / Tables","Empty state shown for new user","1. Login as a brand new employee with no assigned records","List shows No records found message — not a blank white page","","","Medium"),
    # SECTION M — API
    ("M","TC-170","API","GET /quotations with Admin token","1. Get Admin JWT via POST /auth/login\n2. Call GET /quotations","200 OK. Returns all quotations across all enterprise users","","","High"),
    ("M","TC-171","API","GET /quotations with Employee token","1. Get Employee X JWT\n2. Call GET /quotations","200 OK. Returns ONLY Employee X quotations","","","High"),
    ("M","TC-172","API","GET /quotations with Manager token","1. Get Manager A JWT\n2. Call GET /quotations","200 OK. Returns Manager A + all team members quotations","","","High"),
    ("M","TC-173","API","Reject quotation — version does not change","1. Get quotation at v1\n2. Call PATCH /quotations/:id/reject\n3. GET same quotation","Response shows version = 1 (unchanged after rejection)","","","High"),
    ("M","TC-174","API","Revise quotation — version increments","1. Get rejected quotation at v1\n2. Call PATCH /quotations/:id/revise\n3. GET same quotation","Response shows version = 2 after revision","","","High"),
    ("M","TC-175","API","POST /purchase-orders with Employee token","1. Get Employee JWT\n2. Call POST /purchase-orders","403 Forbidden returned","","","High"),
    ("M","TC-176","API","GET /purchase-orders scoped by employee","1. Get Employee X JWT\n2. Call GET /purchase-orders","Returns only POs where Employee X is the assigned owner","","","High"),
    ("M","TC-177","API","Super Admin endpoint with regular JWT","1. Login as Enterprise Admin\n2. Call GET /super-admin/enterprises with enterprise JWT","403 Forbidden — SuperAdminGuard blocks non-super-admin tokens","","","High"),
    ("M","TC-178","API","Reseller endpoint with Super Admin JWT","1. Login as Super Admin\n2. Call GET /resellers/me/tenants with super admin JWT","403 Forbidden — ResellerGuard blocks non-reseller tokens","","","High"),
    ("M","TC-179","API","All endpoints — missing Authorization header","1. Call any protected endpoint without Bearer token","401 Unauthorized response","","","High"),
    ("M","TC-180","API","All endpoints — invalid JWT token","1. Call endpoint with Authorization: Bearer invalidtokenxyz","401 Unauthorized response","","","High"),
    # SECTION N — Negative / Edge Cases
    ("N","TC-181","Negative","Quotation with negative unit price","1. Create quotation\n2. Enter price = -500 on a line item\n3. Submit","Validation error: Price must be a positive value","","","Medium"),
    ("N","TC-182","Negative","PO created without selecting vendor","1. Start PO creation\n2. Leave vendor field blank\n3. Save","Validation error: Vendor is required","","","High"),
    ("N","TC-183","Negative","Invoice payment of 0","1. Open an invoice\n2. Record payment amount = 0","Validation error: Payment amount must be greater than 0","","","Medium"),
    ("N","TC-184","Negative","Edit approved quotation via direct API call","1. Get an Approved quotation ID\n2. Call PATCH /quotations/:id directly","400 or 403 error — must use /revise endpoint to modify","","","High"),
    ("N","TC-185","Negative","Concurrent approval of same quotation by two managers","1. Manager A and Manager B both open same Submitted quotation\n2. Both click Approve simultaneously","Only one approval registered. No duplicate status or conflicting state","","","Medium"),
]

HEADERS = ["TC ID", "Module", "Test Scenario", "Test Steps", "Expected Result", "Actual Result", "Status", "Priority"]
COL_W   = [0.75, 1.1, 2.3, 3.4, 3.0, 1.6, 0.7, 0.75]

current_section = None
for tc in test_cases:
    sec, tc_id, module, scenario, steps, expected, actual, status, priority = tc

    if sec != current_section:
        current_section = sec
        p = doc.add_paragraph()
        run = p.add_run("  " + SECTIONS[sec])
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WHITE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_to_rgb_str(C_SEC_BG))
        pPr.append(shd)

        table = doc.add_table(rows=1, cols=len(HEADERS))
        table.style = 'Table Grid'
        hrow = table.rows[0]
        for i, (h, w) in enumerate(zip(HEADERS, COL_W)):
            cell = hrow.cells[i]
            cell.width = Inches(w)
            shade_cell(cell, C_HDR_BG)
            set_cell_border(cell)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(h)
            cr.bold = True
            cr.font.size = Pt(8)
            cr.font.color.rgb = C_WHITE
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        current_table = table
        row_count = 0

    tr = current_table.add_row()
    row_bg = C_WHITE if row_count % 2 == 0 else C_ALT
    row_count += 1
    vals = [tc_id, module, scenario, steps, expected, actual, status, priority]
    for ci, (val, w) in enumerate(zip(vals, COL_W)):
        cell = tr.cells[ci]
        cell.width = Inches(w)
        if ci == 7 and priority in PRIORITY_RGB:
            shade_cell(cell, PRIORITY_RGB[priority])
        elif ci == 0:
            shade_cell(cell, C_TCID)
        else:
            shade_cell(cell, row_bg)
        set_cell_border(cell)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 6, 7) else WD_ALIGN_PARAGRAPH.LEFT
        cr = cp.add_run(str(val))
        cr.font.size = Pt(8)
        cr.font.color.rgb = C_BODY
        if ci == 0:
            cr.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

doc.add_paragraph()

# ── Coverage Summary ──────────────────────────────────────────────────────────
add_h2(doc, "Test Coverage Summary")
make_table(doc,
    ["Module / Section", "Total TCs", "High", "Medium", "Low"],
    [
        ("Authentication — All Roles", 18, 16, 2, 0),
        ("Super Admin Module", 22, 18, 4, 0),
        ("Reseller Admin Module", 18, 13, 4, 1),
        ("Role-Based Access Control", 9, 9, 0, 0),
        ("Manager Hierarchy", 6, 5, 1, 0),
        ("RFQ Module", 13, 7, 4, 2),
        ("Quotation Module", 24, 17, 6, 1),
        ("Purchase Order Module", 15, 11, 3, 1),
        ("Invoice Module", 14, 10, 4, 0),
        ("Sales Module", 8, 6, 2, 0),
        ("Employee Management", 10, 7, 3, 0),
        ("UI / Navigation / Forms", 12, 4, 5, 3),
        ("API Testing", 11, 11, 0, 0),
        ("Negative / Edge Cases", 5, 2, 3, 0),
        ("TOTAL", 185, 136, 41, 8),
    ],
    [2.8, 0.9, 0.7, 0.8, 0.7]
)

add_note(doc, "Priority: HIGH = business-critical, blocks release if failed  |  MEDIUM = important, should pass before release  |  LOW = minor / cosmetic, can be deferred")
add_note(doc, "v2.0 — Updated to include Super Admin and Reseller Admin roles throughout. All 5 role types covered.")
add_note(doc, "© 2026 VAB Informatics Private Limited — Confidential")

out = r"C:\Users\UNITECH2\Desktop\VABERP_QA_Handover_Document.docx"
doc.save(out)
print(f"Saved: {out}")
