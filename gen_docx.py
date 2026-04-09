from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# ── Colour palette ────────────────────────────────────────────────────────────
C_DARK    = RGBColor(0x1E, 0x29, 0x3B)   # slate-900
C_MID     = RGBColor(0x33, 0x41, 0x55)   # slate-700
C_ACCENT  = RGBColor(0x7C, 0x3A, 0xED)   # violet
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_BODY    = RGBColor(0x1E, 0x29, 0x3B)
C_MUTED   = RGBColor(0x47, 0x55, 0x69)
C_HIGH    = RGBColor(0xFE, 0xE2, 0xE2)   # red-100
C_MED_C   = RGBColor(0xFE, 0xF9, 0xC3)  # yellow-100
C_LOW     = RGBColor(0xDC, 0xFC, 0xE7)   # green-100
C_HDR_BG  = RGBColor(0x1E, 0x29, 0x3B)
C_SEC_BG  = RGBColor(0x33, 0x41, 0x55)
C_ALT     = RGBColor(0xF8, 0xFA, 0xFC)
C_TCID    = RGBColor(0xEF, 0xF6, 0xFF)

def hex_to_rgb_str(c: RGBColor): return '{:02X}{:02X}{:02X}'.format(c[0], c[1], c[2])

# ── Cell shading helper ───────────────────────────────────────────────────────
def shade_cell(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_to_rgb_str(color))
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

# ── Cell border helper ────────────────────────────────────────────────────────
def set_cell_border(cell, color="CBD5E1", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

# ── Paragraph style helpers ───────────────────────────────────────────────────
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = C_DARK
    p.paragraph_format.space_after = Pt(4)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = C_MUTED
    p.paragraph_format.space_after = Pt(16)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '7C3AED')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_ACCENT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_MID
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10)
    run.font.color.rgb = color or C_BODY
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = C_BODY
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_MUTED
    p.paragraph_format.space_after = Pt(6)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(10)

# ── Table helpers ─────────────────────────────────────────────────────────────
def make_table(doc, headers, rows, col_widths, header_bg=C_HDR_BG, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hrow.cells[i]
        cell.width = Inches(w)
        shade_cell(cell, header_bg)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.add_row()
        bg = C_WHITE if ri % 2 == 0 else C_ALT
        for ci, (val, w) in enumerate(zip(row, col_widths)):
            cell = tr.cells[ci]
            cell.width = Inches(w)
            shade_cell(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.color.rgb = C_BODY
            if ci == 0:
                run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_title(doc, "VABERP")
add_subtitle(doc, "QA Handover Document & Test Cases")
add_subtitle(doc, "Version 1.0  ·  2026-04-07  ·  Prepared by: Senior QA Lead")
add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  PART 1 — QA HANDOVER DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "PART 1 — QA HANDOVER DOCUMENT")

# ── Section 1: Project Overview ───────────────────────────────────────────────
add_h2(doc, "1. Project Overview")
make_table(doc,
    ["Field", "Details"],
    [
        ("Product Name", "VABERP — VAB Enterprise ERP"),
        ("Platform", "Web-based (Next.js 14 Frontend · NestJS Backend · PostgreSQL)"),
        ("App URL", "http://64.235.43.187:2262"),
        ("API URL", "http://64.235.43.187:2261"),
        ("Purpose", "End-to-end manufacturing operations management"),
        ("Primary Users", "Admin, Manager, Employee"),
        ("Core Workflow", "RFQ → Quotation → Approval → Purchase Order → Invoice"),
        ("Current Status", "Active / In QA Handover"),
    ],
    [1.5, 5.0]
)
add_body(doc, "VABERP is a manufacturing ERP system that manages the full business cycle: customer enquiries, quotation generation with version control, sales order creation, procurement via RFQ/PO, goods receipts, job card tracking, dispatch, invoicing, and payment recording.")

# ── Section 2: User Roles ─────────────────────────────────────────────────────
add_h2(doc, "2. User Roles & Permissions")
add_h3(doc, "Role Hierarchy")
add_body(doc, "Admin  →  Manager  →  Employee")
add_note(doc, "Each Manager has direct report Employees. Admin oversees all.")

add_h3(doc, "Permission Matrix")
make_table(doc,
    ["Permission", "Admin", "Manager", "Employee"],
    [
        ("View ALL records (all users)", "✅", "❌", "❌"),
        ("View own team's records", "✅", "✅ own team only", "❌"),
        ("View own assigned records", "✅", "✅", "✅"),
        ("Create RFQ / Quotation", "✅", "✅", "✅"),
        ("Submit Quotation for Approval", "✅", "✅", "✅"),
        ("Approve / Reject Quotation", "✅", "✅", "❌"),
        ("Create Purchase Order", "✅", "✅", "❌"),
        ("Generate Invoice", "✅", "✅", "❌"),
        ("Record Payments", "✅", "✅", "❌"),
        ("Manage Users / Roles", "✅", "❌", "❌"),
        ("Access System Settings", "✅", "Limited", "❌"),
        ("View Reports", "✅", "✅ team only", "❌"),
    ],
    [3.2, 0.9, 1.1, 1.1]
)

add_h3(doc, "Key Role Rules")
for rule in [
    "Admin has unrestricted visibility and action access across all modules.",
    "Manager sees their own records AND all records of their direct report employees. Cannot see other managers' team data.",
    "Employee sees only records directly assigned to them.",
    "Role assignment is performed exclusively by Admin.",
    "Hierarchy changes (reassigning employees) take effect immediately.",
]:
    add_bullet(doc, rule)

# ── Section 3: Module Breakdown ───────────────────────────────────────────────
add_h2(doc, "3. Module Breakdown")
modules = [
    ("3.1 RFQ — Request for Quotation", "Allows procurement team to request pricing from vendors. An RFQ can be sent to multiple vendors. Once a vendor responds, the RFQ is converted into a formal Quotation.\nKey states: Draft → Sent → Received → Converted"),
    ("3.2 Quotation", "Core commercial document built with line items (product, qty, unit price, discount). Supports full version control — version increments ONLY on user-initiated revision, never on rejection.\nKey states: Draft → Submitted → Approved | Rejected → Revised (new version)"),
    ("3.3 Purchase Order (PO)", "Generated from an Approved Quotation. Represents formal vendor commitment. Visibility strictly controlled — employees see only POs assigned to them.\nKey states: Draft → Ordered → Received → Invoiced"),
    ("3.4 Invoice", "Generated from a completed PO or Sales Order. Supports partial and full payment recording. Once fully paid, invoice becomes read-only.\nKey states: Unpaid → Partially Paid → Paid"),
    ("3.5 Sales Module", "Manages customer-facing orders from enquiry through production to delivery. Linked to manufacturing job cards and dispatch records.\nKey states: New → In Progress → Dispatched → Completed | Cancelled"),
    ("3.6 Employee & Role Management", "Admin creates and manages user accounts, assigns roles, and sets the reporting structure. Changes to hierarchy are reflected immediately across all modules."),
]
for title, desc in modules:
    add_h3(doc, title)
    for line in desc.split('\n'):
        if line.strip():
            add_body(doc, line)

# ── Section 4: Workflows ──────────────────────────────────────────────────────
add_h2(doc, "4. End-to-End Workflows")

add_h3(doc, "Workflow 1: Procurement Flow — RFQ → PO → Invoice")
steps = [
    "User raises RFQ → selects vendor, adds line items, submits",
    "Vendor responds → RFQ status = Received",
    "User converts RFQ to Quotation → Quotation v1 created, items pre-filled",
    "User submits Quotation for approval → status = Submitted",
    "[Manager/Admin rejects] → Status = Rejected, VERSION STAYS SAME",
    "[User revises] → New version created (v1 → v2), resubmitted",
    "Manager/Admin approves → Status = Approved",
    "User creates Purchase Order from Approved Quotation",
    "Goods received → PO status updated to Received",
    "Invoice generated from PO → status = Unpaid",
    "Payment recorded → Partially Paid → Paid",
]
for i, s in enumerate(steps, 1):
    add_bullet(doc, f"Step {i}: {s}")

add_h3(doc, "Workflow 2: Sales Flow — Enquiry → Invoice")
steps2 = [
    "Customer enquiry captured in system",
    "Quotation created for customer with line items",
    "Quotation approved by Manager/Admin",
    "Sales Order created from Approved Quotation",
    "Job Card generated for manufacturing team",
    "Production stages completed",
    "Goods dispatched → dispatch record created",
    "Invoice raised to customer",
    "Payment received → Invoice marked Paid",
]
for i, s in enumerate(steps2, 1):
    add_bullet(doc, f"Step {i}: {s}")

add_h3(doc, "Workflow 3: Quotation Version Control")
add_body(doc, "Create Quotation → v1 (Draft) → Submit → v1 (Submitted)")
add_bullet(doc, "Approved → v1 (Approved) — version stays 1")
add_bullet(doc, "Rejected → v1 (Rejected) — version STAYS 1, NOT incremented")
add_bullet(doc, "User Revises → v2 (Draft) → Submit → v2 (Submitted) → Approved/Rejected")
add_note(doc, "★ CRITICAL: Version increments ONLY when a user submits a revision — NEVER on rejection alone.")

add_h3(doc, "Workflow 4: Manager Hierarchy Visibility")
add_bullet(doc, "Admin → Sees ALL records from ALL users across ALL modules")
add_bullet(doc, "Manager A → Sees own records + Employee X + Employee Y (direct reports only)")
add_bullet(doc, "Manager B → Sees own records + Employee Z. Cannot see Manager A's team.")
add_bullet(doc, "Employee X → Sees only records assigned to Employee X")

# ── Section 5: Feature Expectations ──────────────────────────────────────────
add_h2(doc, "5. Feature Expectations")
make_table(doc,
    ["Feature", "Expected System Behavior"],
    [
        ("Quotation Version Control", "Version increments ONLY on user revision. Rejection by approver NEVER increments version."),
        ("RFQ to Quotation Conversion", "Converts to Quotation v1 with all line items, quantities, vendor details pre-filled."),
        ("PO Creation Gate", "PO can only be created from Approved quotations. Draft/Submitted/Rejected block PO creation."),
        ("PO Visibility — Employee", "Employee sees ONLY POs where they are the explicitly assigned owner."),
        ("PO Visibility — Manager", "Manager sees all POs belonging to themselves and their direct report employees."),
        ("Invoice Auto-Population", "Invoice amount is pre-populated from PO total. No manual entry required."),
        ("Invoice Status Flow", "Unpaid → Partially Paid → Paid. No reverse transitions allowed."),
        ("Paid Invoice Lock", "Once Paid, all invoice fields are read-only. No further payment edits."),
        ("Duplicate PO Prevention", "Only one PO per approved Quotation. Second attempt must be blocked or warned."),
        ("Duplicate Invoice Prevention", "Only one Invoice per PO. Second attempt must be blocked."),
        ("RBAC Enforcement", "Restricted pages/API endpoints must return 403 or redirect. No data exposed via URL manipulation."),
        ("Manager Cross-Visibility Block", "Manager must never see another manager's team records, even by direct URL/ID access."),
        ("Discount Validation", "Line item discount: 0% minimum, 100% maximum. Values outside this range rejected."),
        ("Quantity Validation", "All quantity fields must be positive integers. Zero and negative values rejected."),
        ("Approved Quotation Edit Lock", "Approved quotation cannot be directly edited. Must initiate Revise action for new version."),
    ],
    [2.0, 4.5]
)

# ── Section 6: Known Issues ───────────────────────────────────────────────────
add_h2(doc, "6. Known Issues / Current Bugs")
make_table(doc,
    ["#", "Module", "Issue Description", "Severity", "Status"],
    [
        ("BUG-001", "Quotation", "[PLACEHOLDER — e.g., Version increments on rejection instead of staying same]", "High", "Open"),
        ("BUG-002", "Purchase Order", "[PLACEHOLDER — e.g., Employee can view POs not assigned to them]", "High", "Open"),
        ("BUG-003", "Invoice", "[PLACEHOLDER — e.g., Invoice amount not auto-populated from PO total]", "Medium", "Open"),
        ("BUG-004", "RBAC", "[PLACEHOLDER — e.g., Manager can see data from other manager's team]", "High", "Open"),
        ("BUG-005", "UI", "[PLACEHOLDER — e.g., Double-click on submit creates duplicate records]", "Medium", "Open"),
        ("BUG-006", "RFQ", "[PLACEHOLDER — fill in during testing]", "—", "Open"),
        ("BUG-007", "Quotation", "[PLACEHOLDER — fill in during testing]", "—", "Open"),
        ("BUG-008", "General", "[PLACEHOLDER — fill in during testing]", "—", "Open"),
    ],
    [0.7, 1.2, 3.3, 0.8, 0.7]
)
add_note(doc, "QA Analyst: Replace placeholders with actual defects found during testing. Log each defect with: module, steps to reproduce, expected vs actual result, severity, screenshot.")

# ── Section 7: Test Environment ───────────────────────────────────────────────
add_h2(doc, "7. Test Environment Setup")
make_table(doc,
    ["Parameter", "Value"],
    [
        ("Application URL", "http://64.235.43.187:2262"),
        ("API Base URL", "http://64.235.43.187:2261"),
        ("Admin Credentials", "Email: [ADMIN_EMAIL]  ·  Password: [ADMIN_PASSWORD]"),
        ("Manager A Credentials", "Email: [MANAGER_A_EMAIL]  ·  Password: [PASSWORD]"),
        ("Manager B Credentials", "Email: [MANAGER_B_EMAIL]  ·  Password: [PASSWORD]"),
        ("Employee X (under Mgr A)", "Email: [EMP_X_EMAIL]  ·  Password: [PASSWORD]"),
        ("Employee Y (under Mgr A)", "Email: [EMP_Y_EMAIL]  ·  Password: [PASSWORD]"),
        ("Employee Z (under Mgr B)", "Email: [EMP_Z_EMAIL]  ·  Password: [PASSWORD]"),
        ("Recommended Browser", "Google Chrome (latest)"),
        ("Secondary Browsers", "Firefox, Microsoft Edge"),
        ("Postman Collection", "[POSTMAN_COLLECTION_LINK]"),
    ],
    [2.0, 4.5]
)

add_h3(doc, "Pre-Test Checklist")
for item in [
    "All 6 user accounts created and accessible",
    "Manager A has Employee X and Employee Y assigned",
    "Manager B has Employee Z assigned",
    "At least 3 vendors created in the system",
    "At least 3 customers created in the system",
    "At least 5 product/item records created",
    "Postman installed with environment configured",
    "Browser dev tools available for console error monitoring",
]:
    add_bullet(doc, item)

# ── Section 8: Test Data ──────────────────────────────────────────────────────
add_h2(doc, "8. Test Data Requirements")
make_table(doc,
    ["Data Category", "Minimum Required", "Notes"],
    [
        ("User Accounts", "1 Admin, 2 Managers, 3 Employees", "Employees split across 2 managers (2+1)"),
        ("Vendors", "3 vendors", "With name, contact, and address"),
        ("Customers", "3 customers", "With name and billing address"),
        ("Products / Line Items", "5–10 products", "With unit prices defined"),
        ("RFQs", "5 RFQs", "Mix of Draft, Sent, Received, Converted states"),
        ("Quotations", "8 quotations", "Mix of v1, v2, Draft, Submitted, Approved, Rejected"),
        ("Purchase Orders", "4 POs", "Assigned: 2 to Emp X, 1 to Emp Y, 1 to Emp Z"),
        ("Invoices", "3 invoices", "Mix of Unpaid, Partially Paid, Paid statuses"),
        ("Sales Orders", "3 sales orders", "Assigned to different employees"),
    ],
    [1.8, 1.9, 2.8]
)

# ── Section 9: API Testing ────────────────────────────────────────────────────
add_h2(doc, "9. API Testing (Postman)")
add_h3(doc, "Setup")
for b in [
    "Base URL: http://64.235.43.187:2261",
    "All endpoints require: Authorization: Bearer <token> header",
    "Obtain token via: POST /auth/login with { email, password }",
    "Create 3 Postman environments: Admin, Manager, Employee (each with own token)",
]:
    add_bullet(doc, b)

add_h3(doc, "Key Endpoints to Test")
make_table(doc,
    ["Method", "Endpoint", "What to Verify"],
    [
        ("POST", "/auth/login", "Role is returned in response. Token issued."),
        ("GET", "/quotations", "Admin = all, Manager = team, Employee = own only"),
        ("POST", "/quotations", "Creates at v1. Returns quotation ID"),
        ("PATCH", "/quotations/:id/submit", "Status → Submitted"),
        ("PATCH", "/quotations/:id/approve", "Status → Approved. Version unchanged"),
        ("PATCH", "/quotations/:id/reject", "Status → Rejected. Version must NOT change"),
        ("PATCH", "/quotations/:id/revise", "Status → Draft. Version must increment"),
        ("POST", "/purchase-orders", "Only Manager/Admin token should succeed (403 for Employee)"),
        ("GET", "/purchase-orders", "Employee sees only assigned. Manager sees team"),
        ("POST", "/invoices", "Blocked if PO not in Received state"),
        ("PATCH", "/invoices/:id/payment", "Records payment, updates status correctly"),
        ("POST", "/rfq/:id/convert", "Creates quotation with pre-filled items"),
    ],
    [0.7, 1.8, 4.0]
)

# ── Section 10: QA Scope ──────────────────────────────────────────────────────
add_h2(doc, "10. QA Scope")
add_h3(doc, "In Scope")
for item in [
    "Functional testing of all 6 modules",
    "Role-based access control — all 3 roles tested independently",
    "Manager hierarchy visibility — cross-team data isolation verified",
    "Quotation version control logic (rejection vs revision behavior)",
    "End-to-end workflow: RFQ → Quotation → PO → Invoice",
    "End-to-end workflow: Enquiry → Sales Order → Invoice",
    "Negative testing (invalid inputs, unauthorized access, duplicate creation)",
    "Boundary testing (0%, 100%, 101% discount; zero/negative quantities)",
    "UI behavior (form validation, button states, double-click prevention, navigation)",
    "API endpoint testing via Postman (auth, scoping, status transitions)",
    "Cross-browser testing (Chrome primary, Firefox and Edge secondary)",
]:
    add_bullet(doc, "✅ " + item)

add_h3(doc, "High Priority Focus Areas")
for item in [
    "Quotation version control — rejection must NOT bump version",
    "PO visibility scoping by role and assignment",
    "RBAC enforcement — both UI and API layer",
    "Manager hierarchy isolation — no cross-team data leakage",
    "Invoice creation gates — only from Received PO",
    "Duplicate record prevention",
]:
    add_bullet(doc, "★ " + item)

# ── Section 11: Out of Scope ──────────────────────────────────────────────────
add_h2(doc, "11. Out of Scope")
for item in [
    "Performance / load testing",
    "Stress testing or concurrent user simulation",
    "Mobile native app (web-responsive only in scope)",
    "Email notification content/delivery testing",
    "Third-party integrations",
    "Database backup and recovery testing",
    "Infrastructure / server configuration testing",
    "Accessibility (WCAG) compliance testing",
]:
    add_bullet(doc, "❌ " + item)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
#  PART 2 — TEST CASES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, "PART 2 — TEST CASES  (145 Total)")
add_note(doc, "Columns: TC ID | Module | Test Scenario | Test Steps | Expected Result | Actual Result | Status | Priority")

SECTIONS = {
    "A": "Section A — Authentication (TC-001 to TC-013)",
    "B": "Section B — Role-Based Access Control (TC-014 to TC-022)",
    "C": "Section C — Manager Hierarchy & Visibility (TC-023 to TC-028)",
    "D": "Section D — RFQ Module (TC-029 to TC-041)",
    "E": "Section E — Quotation Module (TC-042 to TC-065)",
    "F": "Section F — Purchase Order Module (TC-066 to TC-080)",
    "G": "Section G — Invoice Module (TC-081 to TC-094)",
    "H": "Section H — Sales Module (TC-095 to TC-102)",
    "I": "Section I — Employee & Role Management (TC-103 to TC-112)",
    "J": "Section J — UI, Navigation & Forms (TC-113 to TC-124)",
    "K": "Section K — API Testing (TC-125 to TC-135)",
    "L": "Section L — Negative & Edge Cases (TC-136 to TC-145)",
}

PRIORITY_RGB = {
    "High":   RGBColor(0xFE, 0xE2, 0xE2),
    "Medium": RGBColor(0xFE, 0xF9, 0xC3),
    "Low":    RGBColor(0xDC, 0xFC, 0xE7),
}

test_cases = [
    ("A","TC-001","Auth","Admin login with valid credentials","1. Go to /login\n2. Enter admin email & password\n3. Click Login","Admin dashboard loads with full navigation menu visible","","","High"),
    ("A","TC-002","Auth","Manager login with valid credentials","1. Go to /login\n2. Enter manager credentials\n3. Click Login","Manager dashboard loads; team-scoped menu items visible","","","High"),
    ("A","TC-003","Auth","Employee login with valid credentials","1. Go to /login\n2. Enter employee credentials\n3. Click Login","Employee dashboard loads; restricted menu (no PO create, no approvals)","","","High"),
    ("A","TC-004","Auth","Login with invalid password","1. Enter valid email + wrong password\n2. Click Login","Error: Invalid credentials. No redirect. Token not issued","","","High"),
    ("A","TC-005","Auth","Login with unregistered email","1. Enter email not in system\n2. Enter any password\n3. Click Login","Error: User not found or Invalid credentials","","","High"),
    ("A","TC-006","Auth","Login with blank email field","1. Leave email blank\n2. Fill password\n3. Click Login","Frontend validation: Email is required","","","Medium"),
    ("A","TC-007","Auth","Login with blank password field","1. Fill email\n2. Leave password blank\n3. Click Login","Frontend validation: Password is required","","","Medium"),
    ("A","TC-008","Auth","Login with both fields blank","1. Leave both fields blank\n2. Click Login","Both field validation errors shown. No API call made","","","Medium"),
    ("A","TC-009","Auth","Session persistence after page refresh","1. Login as any user\n2. Refresh browser","User remains logged in. Dashboard reloads without re-authentication","","","High"),
    ("A","TC-010","Auth","Logout clears session","1. Login as any user\n2. Click Logout","Session cleared. Redirected to /login. Back button does not return to dashboard","","","High"),
    ("A","TC-011","Auth","Expired session redirects to login","1. Login\n2. Manually clear auth token from browser storage\n3. Try to navigate","Redirected to /login. Unauthorized message shown","","","High"),
    ("A","TC-012","Auth","Brute force — repeated wrong password","1. Enter wrong password 5 times consecutively","Account locked or rate-limit error message shown","","","Medium"),
    ("A","TC-013","Auth","SQL injection in login email field","1. Enter ' OR 1=1 -- in email field\n2. Click Login","Login rejected. No data exposed. Standard error message shown","","","High"),
    ("B","TC-014","RBAC","Employee cannot access PO creation","1. Login as Employee\n2. Navigate to /purchase-orders/new","Access denied or button not visible. No PO form rendered","","","High"),
    ("B","TC-015","RBAC","Employee cannot generate invoices","1. Login as Employee\n2. Open a Received PO\n3. Check for Generate Invoice action","Generate Invoice button not visible to Employee","","","High"),
    ("B","TC-016","RBAC","Employee cannot access User Management","1. Login as Employee\n2. Navigate to /settings/users via URL","Redirect to dashboard or 403 error. No user list shown","","","High"),
    ("B","TC-017","RBAC","Manager cannot access User Management","1. Login as Manager\n2. Navigate to /settings/users","Access denied. Create/Edit user actions unavailable","","","High"),
    ("B","TC-018","RBAC","Manager can approve quotations","1. Login as Manager\n2. Open Submitted quotation\n3. Click Approve","Approval action succeeds. Status changes to Approved","","","High"),
    ("B","TC-019","RBAC","Employee cannot approve quotations","1. Login as Employee\n2. Open a Submitted quotation","Approve button not visible. Cannot trigger approval action","","","High"),
    ("B","TC-020","RBAC","Admin can access all modules","1. Login as Admin\n2. Navigate through every menu item","All modules accessible without restriction","","","High"),
    ("B","TC-021","RBAC","Direct URL access by unauthorized role","1. Login as Employee\n2. Manually type /purchase-orders/new in URL bar","Redirected away or 403 error shown. No form rendered","","","High"),
    ("B","TC-022","RBAC","API call with Employee token to Manager endpoint","1. Login as Employee via Postman\n2. Call POST /purchase-orders with employee JWT","403 Forbidden response returned","","","High"),
    ("C","TC-023","Hierarchy","Manager A sees Employee X and Y records","1. Employee X and Y (both under Manager A) each create a quotation\n2. Login as Manager A\n3. Open Quotations list","Records of both Employee X and Employee Y visible to Manager A","","","High"),
    ("C","TC-024","Hierarchy","Manager A cannot see Employee Z records (under B)","1. Employee Z (under Manager B) creates a quotation\n2. Login as Manager A\n3. Open Quotations list","Employee Z quotation NOT in Manager A list","","","High"),
    ("C","TC-025","Hierarchy","Manager B cannot see Manager A team records","1. Employee X (under Manager A) creates a record\n2. Login as Manager B","Employee X record not visible to Manager B","","","High"),
    ("C","TC-026","Hierarchy","Admin sees all records from all users","1. Employees under Manager A and B each create records\n2. Login as Admin","All records from all users visible","","","High"),
    ("C","TC-027","Hierarchy","Employee reassigned — visibility updates immediately","1. Employee X under Manager A\n2. Admin reassigns to Manager B\n3. Login as Manager B","Employee X records visible to Manager B. No longer visible to Manager A","","","High"),
    ("C","TC-028","Hierarchy","Manager own records visible to themselves","1. Manager A creates a quotation directly\n2. Login as Manager A","Manager A own quotation visible in their list","","","Medium"),
    ("D","TC-029","RFQ","Create RFQ with valid data","1. Login\n2. Navigate to RFQ\n3. Fill vendor, items, quantities\n4. Submit","RFQ created with status Draft/Sent and unique reference number","","","High"),
    ("D","TC-030","RFQ","Create RFQ with missing vendor","1. Fill RFQ items\n2. Leave vendor blank\n3. Submit","Validation error: Vendor is required","","","High"),
    ("D","TC-031","RFQ","Create RFQ with zero quantity","1. Create RFQ\n2. Set item qty = 0\n3. Submit","Validation error: Quantity must be greater than 0","","","Medium"),
    ("D","TC-032","RFQ","Create RFQ with negative quantity","1. Create RFQ\n2. Enter qty = -5\n3. Submit","Validation error: negative values not allowed","","","Medium"),
    ("D","TC-033","RFQ","Create RFQ with no line items","1. Fill vendor\n2. Add no items\n3. Submit","Validation error: At least one line item is required","","","High"),
    ("D","TC-034","RFQ","Convert RFQ to Quotation","1. Open a submitted/received RFQ\n2. Click Convert to Quotation","Quotation created at v1 with all items pre-filled from RFQ","","","High"),
    ("D","TC-035","RFQ","Quotation from RFQ pre-fills all line items","1. Create RFQ with 3 items\n2. Convert to Quotation\n3. Open the quotation","All 3 items present in quotation with matching quantity and description","","","High"),
    ("D","TC-036","RFQ","Employee sees only their own RFQs","1. Employee X creates an RFQ\n2. Login as Employee Y","Employee Y cannot see Employee X RFQ","","","High"),
    ("D","TC-037","RFQ","Manager sees all team RFQs","1. Employee X (under Manager A) creates an RFQ\n2. Login as Manager A","RFQ appears in Manager A list","","","High"),
    ("D","TC-038","RFQ","Duplicate RFQ same vendor same reference","1. Create RFQ for Vendor A ref #001\n2. Create another with same vendor and reference","System warns or blocks exact duplicate","","","Medium"),
    ("D","TC-039","RFQ","RFQ list displays correct status badges","1. Create RFQs in different states\n2. Open RFQ list","Status badges correctly show Draft, Sent, Received per record","","","Low"),
    ("D","TC-040","RFQ","Delete a Draft RFQ","1. Create an RFQ\n2. Keep it in Draft\n3. Delete it","RFQ removed from list. No linked records affected","","","Medium"),
    ("D","TC-041","RFQ","Cannot delete a converted RFQ","1. Convert RFQ to Quotation\n2. Try to delete the original RFQ","Deletion blocked: RFQ has a linked Quotation","","","Medium"),
    ("E","TC-042","Quotation","Create quotation with valid line items","1. Navigate to Create Quotation\n2. Add items with qty, unit price, discount\n3. Save","Quotation saved at v1, status = Draft, total calculated correctly","","","High"),
    ("E","TC-043","Quotation","Total calculation — single item with discount","1. Add Item: qty=2, price=1000, discount=10%","Total = (2x1000) - 10% = 1800","","","High"),
    ("E","TC-044","Quotation","Total calculation — multiple items","1. Item A: qty=1, price=500, 0% discount\n2. Item B: qty=2, price=300, 5% discount","Total = 500 + 570 = 1070","","","High"),
    ("E","TC-045","Quotation","Discount exactly 100%","1. Create quotation\n2. Enter discount = 100% on a line item","Accepted — line item cost = 0. Total reflects correctly","","","Medium"),
    ("E","TC-046","Quotation","Discount over 100%","1. Create quotation\n2. Enter discount = 101%\n3. Save","Validation error: Discount cannot exceed 100%","","","High"),
    ("E","TC-047","Quotation","Discount = 0% no discount","1. Leave discount blank or enter 0 on a line item","No discount applied. Full price used in total","","","Low"),
    ("E","TC-048","Quotation","Create quotation with no line items","1. Open Create Quotation\n2. Add no line items\n3. Submit","Validation error: At least one line item is required","","","High"),
    ("E","TC-049","Quotation","Submit quotation for approval","1. Open a Draft quotation\n2. Click Submit","Status = Submitted. Approve/Reject options now available to approvers","","","High"),
    ("E","TC-050","Quotation","Approve a submitted quotation","1. Login as Manager or Admin\n2. Open Submitted quotation\n3. Click Approve","Status = Approved. Version number unchanged. Edit fields locked","","","High"),
    ("E","TC-051","Quotation","Reject a submitted quotation","1. Login as Manager or Admin\n2. Open Submitted quotation\n3. Click Reject\n4. Enter rejection reason","Status = Rejected. Version number stays the same (v1 stays v1)","","","High"),
    ("E","TC-052","Quotation","CRITICAL: Version does NOT increment on rejection","1. Create Quotation (version = 1)\n2. Submit\n3. Manager rejects\n4. Check version field","Version field = 1 unchanged. Rejection must NEVER auto-increment version","","","High"),
    ("E","TC-053","Quotation","Revise a rejected quotation","1. Open Rejected quotation\n2. Click Revise\n3. Modify at least one line item\n4. Resubmit","New version created: v1 to v2. Original v1 preserved in version history","","","High"),
    ("E","TC-054","Quotation","CRITICAL: Version increments on user revision","1. Reject a v1 quotation\n2. User clicks Revise\n3. Check version after save","Version = 2. Re-submission shows v2 to approver","","","High"),
    ("E","TC-055","Quotation","Multiple revision cycles — correct versioning","1. v1 Reject Revise v2 Reject Revise v3","Versions: v1, v2, v3 in correct sequence. All listed in history","","","High"),
    ("E","TC-056","Quotation","Version history shows all revisions with metadata","1. Create quotation with 3 revision cycles\n2. Open Version History tab","v1, v2, v3 listed with timestamp, status, and who approved/rejected","","","Medium"),
    ("E","TC-057","Quotation","Approved quotation fields are read-only","1. Open an Approved quotation\n2. Try to click and edit any field","All fields read-only. Only Revise action button available","","","High"),
    ("E","TC-058","Quotation","Revise approved quotation creates new version","1. Open Approved quotation\n2. Click Revise\n3. Edit and resubmit","New version (e.g. v2) created in Draft. Previous approved version preserved","","","High"),
    ("E","TC-059","Quotation","Cannot re-submit already Submitted quotation","1. Submit a quotation (status = Submitted)\n2. Click Submit again","Submit button disabled or already-submitted message shown","","","Medium"),
    ("E","TC-060","Quotation","Employee cannot approve own submitted quotation","1. Login as Employee\n2. Create and submit a quotation\n3. Check for Approve button","Approve button NOT visible to the Employee who submitted it","","","High"),
    ("E","TC-061","Quotation","Quotation linked to correct customer","1. Create quotation for Customer A\n2. Open quotation detail","Customer A shown in customer field. No data leakage from other customers","","","High"),
    ("E","TC-062","Quotation","Search quotation by customer name","1. Open Quotation list\n2. Search by customer name","Only quotations linked to that customer are shown","","","Medium"),
    ("E","TC-063","Quotation","Filter quotations by status","1. Open Quotation list\n2. Apply filter Approved","Only Approved quotations displayed in list","","","Medium"),
    ("E","TC-064","Quotation","Quotation PDF/print generation","1. Open an Approved quotation\n2. Click Download or Print","PDF generated with correct line items, totals, customer details, and version","","","Medium"),
    ("E","TC-065","Quotation","Duplicate quotation reference prevention","1. Create Quotation ref #Q-001 for Customer A\n2. Attempt to create another with same ref/customer","System warns or prevents exact duplicate quotation creation","","","Medium"),
    ("F","TC-066","Purchase Order","Create PO from approved quotation","1. Open Approved quotation\n2. Click Create PO\n3. Assign to Employee\n4. Save","PO created with correct amounts and assigned to selected employee","","","High"),
    ("F","TC-067","Purchase Order","PO amount matches approved quotation total","1. Approved quotation total = 25000\n2. Create PO from it","PO total = 25000. No discrepancy between quotation and PO","","","High"),
    ("F","TC-068","Purchase Order","Cannot create PO from Draft quotation","1. Open a Draft quotation\n2. Check for Create PO option","Create PO button not visible or disabled","","","High"),
    ("F","TC-069","Purchase Order","Cannot create PO from Rejected quotation","1. Open a Rejected quotation\n2. Check for Create PO option","Create PO button not available for rejected quotations","","","High"),
    ("F","TC-070","Purchase Order","Cannot create PO from Submitted quotation","1. Open a Submitted (pending approval) quotation\n2. Check Create PO","Create PO disabled — must wait for approval","","","High"),
    ("F","TC-071","Purchase Order","Duplicate PO prevention for same quotation","1. Create PO from Approved quotation\n2. Try to create another PO from same quotation","System blocks or warns: A PO already exists for this quotation","","","High"),
    ("F","TC-072","Purchase Order","Employee sees only their assigned POs","1. Create PO assigned to Employee X\n2. Login as Employee X\n3. Open PO list","PO assigned to Employee X is visible","","","High"),
    ("F","TC-073","Purchase Order","Employee cannot see PO assigned to another employee","1. Create PO assigned to Employee Y\n2. Login as Employee X\n3. Open PO list","Employee Y PO is NOT in Employee X list","","","High"),
    ("F","TC-074","Purchase Order","Manager sees all team POs","1. Create POs for Employee X and Y (both under Manager A)\n2. Login as Manager A","Both POs from Employee X and Y visible to Manager A","","","High"),
    ("F","TC-075","Purchase Order","Update PO status to Received","1. Open an active PO\n2. Update status to Received","Status updates to Received. Invoice generation now enabled","","","High"),
    ("F","TC-076","Purchase Order","PO cannot be deleted if linked invoice exists","1. Generate an invoice from PO\n2. Try to delete the PO","Deletion blocked: Cannot delete a PO with an associated invoice","","","High"),
    ("F","TC-077","Purchase Order","PO list sorted by date newest first","1. Create 3 POs on different dates\n2. Open PO list","Most recently created PO appears first","","","Low"),
    ("F","TC-078","Purchase Order","Search PO by vendor name","1. Open PO list\n2. Search Vendor A","Only POs linked to Vendor A shown","","","Medium"),
    ("F","TC-079","Purchase Order","PO reference number is unique","1. Create multiple POs","Each PO has a unique system-generated reference number","","","High"),
    ("F","TC-080","Purchase Order","Reassign PO updates visibility","1. PO assigned to Employee X\n2. Edit PO and reassign to Employee Y","PO now visible to Employee Y. No longer visible to Employee X","","","Medium"),
    ("G","TC-081","Invoice","Generate invoice from Received PO","1. Open PO with status = Received\n2. Click Generate Invoice","Invoice created with amount pre-filled from PO. Status = Unpaid","","","High"),
    ("G","TC-082","Invoice","Invoice amount auto-populated from PO","1. PO total = 40000\n2. Generate invoice","Invoice amount = 40000. No manual entry required","","","High"),
    ("G","TC-083","Invoice","Invoice default status is Unpaid","1. Generate any new invoice","Invoice status = Unpaid on creation","","","High"),
    ("G","TC-084","Invoice","Record partial payment","1. Invoice total = 50000\n2. Record payment of 20000","Status = Partially Paid. Remaining balance = 30000 shown","","","High"),
    ("G","TC-085","Invoice","Record full payment","1. Invoice total = 50000\n2. Record payment of 50000","Status = Paid. Balance = 0. Invoice marked complete","","","High"),
    ("G","TC-086","Invoice","Payment amount exceeds invoice total","1. Invoice = 50000\n2. Attempt to record 60000 payment","Validation error: Payment cannot exceed invoice total","","","High"),
    ("G","TC-087","Invoice","Multiple partial payments until fully paid","1. Invoice = 50000\n2. Record 20000 then 20000 then 10000","Each payment recorded. Final payment triggers status = Paid","","","High"),
    ("G","TC-088","Invoice","Employee cannot generate invoice","1. Login as Employee\n2. Open a Received PO","Generate Invoice button not visible to Employee role","","","High"),
    ("G","TC-089","Invoice","Invoice linked to correct PO with clickable reference","1. Generate invoice from PO #P-001\n2. Open invoice detail","Invoice shows reference to PO #P-001 with a working link","","","Medium"),
    ("G","TC-090","Invoice","Navigate from invoice back to source PO","1. Open an invoice\n2. Click PO reference link","Correctly navigates to the originating PO record","","","Medium"),
    ("G","TC-091","Invoice","Filter invoice list by payment status","1. Open Invoices list\n2. Filter by Paid","Only Paid invoices shown in list","","","Medium"),
    ("G","TC-092","Invoice","Invoice PDF/print generation","1. Open an invoice\n2. Click Download or Print","PDF generated with correct amounts, payment status, and customer details","","","Medium"),
    ("G","TC-093","Invoice","Paid invoice is read-only","1. Open invoice with status = Paid\n2. Try to modify any field or add a payment","All fields and actions locked. No modification possible","","","High"),
    ("G","TC-094","Invoice","Duplicate invoice prevention for same PO","1. Generate invoice from PO #P-001\n2. Try to generate another invoice from same PO","Blocked: An invoice already exists for this Purchase Order","","","High"),
    ("H","TC-095","Sales","Create customer sales order with valid data","1. Navigate to Sales\n2. Create order for Customer A\n3. Add line items\n4. Save","Sales order saved with unique reference and correct status","","","High"),
    ("H","TC-096","Sales","Sales order linked to approved quotation","1. Approve a customer quotation\n2. Convert to Sales Order","Sales order references the originating quotation. Line items match","","","High"),
    ("H","TC-097","Sales","Sales order triggers job card creation","1. Create a Sales Order\n2. Click Create Job Card","Job card created with reference back to the sales order","","","High"),
    ("H","TC-098","Sales","Employee sees only assigned sales orders","1. Create 2 sales orders (1 assigned to X, 1 to Y)\n2. Login as Employee X","Only Employee X sales order visible in their list","","","High"),
    ("H","TC-099","Sales","Manager sees all team sales orders","1. Sales orders assigned to Employee X and Y (both under Manager A)\n2. Login as Manager A","Both sales orders visible to Manager A","","","High"),
    ("H","TC-100","Sales","Update sales order status","1. Open a sales order\n2. Change status from In Progress to Dispatched","Status updates correctly. History log reflects the change with timestamp","","","Medium"),
    ("H","TC-101","Sales","Sales order total calculation","1. Add 2 line items with qty, price, and discount\n2. Check total","Total reflects correct arithmetic with all discounts applied","","","High"),
    ("H","TC-102","Sales","Cancel a sales order","1. Open an In Progress sales order\n2. Click Cancel","Status = Cancelled. Associated job card status also updated to Cancelled","","","Medium"),
    ("I","TC-103","Employee Mgmt","Admin creates new employee user","1. Login as Admin\n2. Navigate to Users\n3. Create user role = Employee\n4. Assign to Manager A","User created. Appears in Manager A team immediately","","","High"),
    ("I","TC-104","Employee Mgmt","Admin creates new manager user","1. Login as Admin\n2. Create user with role = Manager","Manager account created. Can see team records once employees assigned","","","High"),
    ("I","TC-105","Employee Mgmt","Admin assigns employee to a manager","1. Login as Admin\n2. Edit Employee X\n3. Set reporting manager = Manager A","Employee X appears under Manager A team","","","High"),
    ("I","TC-106","Employee Mgmt","Admin reassigns employee to new manager","1. Employee X under Manager A\n2. Admin reassigns to Manager B","Manager B now sees Employee X records. Manager A no longer sees them","","","High"),
    ("I","TC-107","Employee Mgmt","Manager cannot create user accounts","1. Login as Manager\n2. Navigate to User Management","Create/Edit user actions not available to Manager role","","","High"),
    ("I","TC-108","Employee Mgmt","Admin upgrades employee role to Manager","1. Login as Admin\n2. Edit Employee X\n3. Change role to Manager","Employee X gets Manager-level access on next login","","","Medium"),
    ("I","TC-109","Employee Mgmt","Admin deactivates a user account","1. Login as Admin\n2. Deactivate Employee X","Employee X cannot login. Existing records remain visible to Manager/Admin","","","Medium"),
    ("I","TC-110","Employee Mgmt","Deactivated user cannot login","1. Admin deactivates Employee X\n2. Employee X attempts to login with correct credentials","Login rejected: Account deactivated or Invalid credentials","","","High"),
    ("I","TC-111","Employee Mgmt","Admin edits user profile details","1. Login as Admin\n2. Edit name/email of a user\n3. Save","Changes saved. Updated information reflects across the system","","","Medium"),
    ("I","TC-112","Employee Mgmt","Duplicate email on user creation","1. Login as Admin\n2. Create user with email already registered in system","Validation error: Email already registered","","","High"),
    ("J","TC-113","UI / Navigation","Sidebar navigation — all links work","1. Login as Admin\n2. Click each sidebar menu item","Correct page loads for each. No 404 errors encountered","","","High"),
    ("J","TC-114","UI / Navigation","Breadcrumb navigation accuracy","1. Navigate to a nested page (e.g. Quotation detail)\n2. Check breadcrumb","Breadcrumb shows correct path: Home > Quotations > #Q-001","","","Low"),
    ("J","TC-115","UI / Navigation","Back button preserves list filter state","1. Apply a filter on quotation list\n2. Open a record\n3. Press browser back","Returns to list with same filter still applied","","","Low"),
    ("J","TC-116","UI / Forms","Double-click submit prevention","1. Fill a quotation form completely\n2. Rapidly double-click Submit button","Only one record created. No duplicate quotation in the list","","","High"),
    ("J","TC-117","UI / Forms","Required field highlighting on submit","1. Open any create form\n2. Leave required fields blank\n3. Click Submit","All blank required fields highlighted in red with descriptive error text","","","High"),
    ("J","TC-118","UI / Forms","Form data persists on browser tab switch","1. Start filling a form\n2. Switch to another browser tab\n3. Return","Form data still intact. No data lost","","","Medium"),
    ("J","TC-119","UI / Forms","Long text input in description field","1. Enter 1000+ characters in any description field\n2. Save","Content accepted and saved. Displayed correctly on detail view","","","Low"),
    ("J","TC-120","UI / Forms","XSS injection in text fields","1. Enter script tag in a name or description field\n2. Save and reload","Input sanitized. Script does not execute. Stored as literal text","","","High"),
    ("J","TC-121","UI / Tables","Pagination on large dataset","1. Login as Admin\n2. Open list with 50+ records","Data loads in pages. Pagination controls work correctly","","","Medium"),
    ("J","TC-122","UI / Tables","Column sorting works correctly","1. Open any list table\n2. Click a column header","List re-orders ascending then descending on repeated clicks","","","Medium"),
    ("J","TC-123","UI / Tables","Export to Excel or CSV","1. Open any list\n2. Click Export button","File downloaded with correct column headers and all visible data","","","Medium"),
    ("J","TC-124","UI / Tables","Empty state shown for new user","1. Login as a brand new employee with no assigned records","List shows No records found message — not a blank white page","","","Medium"),
    ("K","TC-125","API","GET /quotations with Admin token","1. Get Admin JWT via POST /auth/login\n2. Call GET /quotations","200 OK. Returns all quotations across all users","","","High"),
    ("K","TC-126","API","GET /quotations with Employee token","1. Get Employee X JWT\n2. Call GET /quotations","200 OK. Returns ONLY Employee X quotations. No other user data","","","High"),
    ("K","TC-127","API","GET /quotations with Manager token","1. Get Manager A JWT\n2. Call GET /quotations","200 OK. Returns Manager A + all team members quotations","","","High"),
    ("K","TC-128","API","POST /quotations with missing required field","1. Call POST /quotations without customer_id in body","400 Bad Request with descriptive validation message","","","High"),
    ("K","TC-129","API","Reject quotation — version does not change","1. Get quotation at v1\n2. Call PATCH /quotations/:id/reject\n3. GET same quotation","Response shows version = 1 (unchanged after rejection)","","","High"),
    ("K","TC-130","API","Revise quotation — version increments","1. Get rejected quotation at v1\n2. Call PATCH /quotations/:id/revise\n3. GET same quotation","Response shows version = 2 after revision","","","High"),
    ("K","TC-131","API","POST /purchase-orders with Employee token","1. Get Employee JWT\n2. Call POST /purchase-orders","403 Forbidden returned","","","High"),
    ("K","TC-132","API","GET /purchase-orders scoped by employee","1. Get Employee X JWT\n2. Call GET /purchase-orders","Returns only POs where Employee X is the assigned owner","","","High"),
    ("K","TC-133","API","POST /invoices — duplicate for same PO","1. Generate invoice for PO #P-001\n2. Call POST /invoices again for same PO ID","409 Conflict or 400 validation error returned","","","High"),
    ("K","TC-134","API","All endpoints — missing Authorization header","1. Call any protected endpoint without Bearer token","401 Unauthorized response","","","High"),
    ("K","TC-135","API","All endpoints — invalid JWT token","1. Call endpoint with Authorization: Bearer invalidtokenxyz","401 Unauthorized response","","","High"),
    ("L","TC-136","Negative","Quotation with negative unit price","1. Create quotation\n2. Enter price = -500 on a line item\n3. Submit","Validation error: Price must be a positive value","","","Medium"),
    ("L","TC-137","Negative","PO created without selecting vendor","1. Start PO creation\n2. Leave vendor field blank\n3. Save","Validation error: Vendor is required","","","High"),
    ("L","TC-138","Negative","Invoice payment of 0","1. Open an invoice\n2. Record payment amount = 0","Validation error: Payment amount must be greater than 0","","","Medium"),
    ("L","TC-139","Negative","Navigate to URL of deleted record","1. Delete a quotation\n2. Navigate to its direct URL","404 Not Found or Record not found message shown","","","Medium"),
    ("L","TC-140","Negative","Edit approved quotation via direct API call","1. Get an Approved quotation ID\n2. Call PATCH /quotations/:id directly with field modifications","400 or 403 error — must use /revise endpoint to modify","","","High"),
    ("L","TC-141","Negative","Generate invoice from PO not yet Received","1. Open a PO with status = Ordered (not Received)\n2. Attempt to generate invoice","Invoice generation blocked: PO must be in Received status","","","High"),
    ("L","TC-142","Negative","Delete employee who has open active records","1. Employee X has open quotations and POs\n2. Admin tries to delete Employee X","Deletion blocked or prompt to reassign records shown","","","High"),
    ("L","TC-143","Negative","Assign PO to non-existent employee via API","1. Call POST /purchase-orders with employee_id = 99999","400 or 404 — employee not found error","","","Medium"),
    ("L","TC-144","Negative","XSS in quotation description via API","1. Call POST /quotations with XSS payload in description\n2. GET the record","Script not executed on retrieval. Stored and returned as sanitized plain string","","","High"),
    ("L","TC-145","Negative","Concurrent approval of same quotation by two managers","1. Manager A and Manager B both open same Submitted quotation\n2. Both click Approve simultaneously","Only one approval registered. No duplicate status or conflicting state","","","Medium"),
]

HEADERS = ["TC ID", "Module", "Test Scenario", "Test Steps", "Expected Result", "Actual Result", "Status", "Priority"]
COL_W   = [0.75, 1.1, 2.3, 3.4, 3.0, 1.6, 0.7, 0.75]

current_section = None
for tc in test_cases:
    sec, tc_id, module, scenario, steps, expected, actual, status, priority = tc

    if sec != current_section:
        current_section = sec
        # Section heading paragraph
        p = doc.add_paragraph()
        run = p.add_run("  " + SECTIONS[sec])
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WHITE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        # background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_to_rgb_str(C_SEC_BG))
        pPr.append(shd)

        # Table header
        table = doc.add_table(rows=1, cols=len(HEADERS))
        table.style = 'Table Grid'
        hrow = table.rows[0]
        for i, (h, w) in enumerate(zip(HEADERS, COL_W)):
            cell = hrow.cells[i]
            cell.width = Inches(w)
            shade_cell(cell, C_HDR_BG)
            set_cell_border(cell)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(h)
            cr.bold = True
            cr.font.size = Pt(8)
            cr.font.color.rgb = C_WHITE
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        current_table = table
        row_count = 0

    # Data row
    tr = current_table.add_row()
    row_bg = C_WHITE if row_count % 2 == 0 else C_ALT
    row_count += 1
    vals = [tc_id, module, scenario, steps, expected, actual, status, priority]
    for ci, (val, w) in enumerate(zip(vals, COL_W)):
        cell = tr.cells[ci]
        cell.width = Inches(w)
        if ci == 7 and priority in PRIORITY_RGB:
            shade_cell(cell, PRIORITY_RGB[priority])
        elif ci == 0:
            shade_cell(cell, C_TCID)
        else:
            shade_cell(cell, row_bg)
        set_cell_border(cell)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 6, 7) else WD_ALIGN_PARAGRAPH.LEFT
        cr = cp.add_run(str(val))
        cr.font.size = Pt(8)
        cr.font.color.rgb = C_BODY
        if ci == 0:
            cr.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

doc.add_paragraph()

# ── Coverage Summary ──────────────────────────────────────────────────────────
add_h2(doc, "Test Coverage Summary")
make_table(doc,
    ["Module / Section", "Total TCs", "High", "Medium", "Low"],
    [
        ("Authentication", 13, 9, 3, 1),
        ("Role-Based Access Control", 9, 9, 0, 0),
        ("Manager Hierarchy", 6, 5, 1, 0),
        ("RFQ Module", 13, 7, 4, 2),
        ("Quotation Module", 24, 17, 6, 1),
        ("Purchase Order Module", 15, 11, 3, 1),
        ("Invoice Module", 14, 10, 4, 0),
        ("Sales Module", 8, 6, 2, 0),
        ("Employee Management", 10, 7, 3, 0),
        ("UI / Navigation / Forms", 12, 4, 5, 3),
        ("API Testing", 11, 11, 0, 0),
        ("Negative / Edge Cases", 10, 6, 4, 0),
        ("TOTAL", 145, 102, 35, 8),
    ],
    [2.8, 0.9, 0.7, 0.8, 0.7]
)

add_note(doc, "Priority: HIGH = business-critical, blocks release if failed  |  MEDIUM = important, should pass before release  |  LOW = minor / cosmetic, can be deferred")
add_note(doc, "© 2026 VAB Informatics Private Limited — Confidential")

out = r"C:\Users\UNITECH2\Desktop\VABERP_QA_Handover_Document.docx"
doc.save(out)
print("Saved:", out)
