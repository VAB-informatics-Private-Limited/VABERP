"""Convert docs/ERP_Enterprise.md → ERP_Enterprise.xlsx and ERP_Enterprise.docx.

Run:
    python docs/convert_erp_enterprise.py
"""
from __future__ import annotations
import os
import re
from dataclasses import dataclass, field
from typing import List, Optional

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "ERP_Enterprise.md")
XLSX = os.path.join(HERE, "ERP_Enterprise.xlsx")
DOCX = os.path.join(HERE, "ERP_Enterprise.docx")


# ─── Markdown parsing into a flat block list ─────────────────────────────────
@dataclass
class Block:
    kind: str            # 'h1','h2','h3','h4','para','ul','ol','table','code','hr','blockquote'
    text: str = ""       # heading or paragraph text
    items: List[str] = field(default_factory=list)   # for ul / ol / blockquote / code
    headers: List[str] = field(default_factory=list) # for table
    rows: List[List[str]] = field(default_factory=list)  # for table
    section: int = 0     # H2 index (0 before any H2)


def strip_inline(s: str) -> str:
    """Remove markdown inline emphasis but keep readable text."""
    s = s.replace("\r", "")
    # inline code
    s = re.sub(r"`([^`]+)`", r"\1", s)
    # bold/italic
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    s = re.sub(r"__([^_]+)__", r"\1", s)
    s = re.sub(r"_([^_]+)_", r"\1", s)
    # links [text](url) → text
    s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", s)
    # html breaks → newlines (used in TC tables)
    s = s.replace("<br/>", "\n").replace("<br>", "\n")
    return s.strip()


def parse_table_row(line: str) -> List[str]:
    # Split on `|` but respect leading/trailing pipes
    parts = line.strip()
    if parts.startswith("|"):
        parts = parts[1:]
    if parts.endswith("|"):
        parts = parts[:-1]
    cells = [strip_inline(c.strip()) for c in parts.split("|")]
    return cells


def parse_markdown(path: str) -> List[Block]:
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    blocks: List[Block] = []
    i = 0
    section_idx = 0
    in_code = False
    code_buf: List[str] = []

    def flush_paragraph(buf: List[str]):
        text = " ".join(line.strip() for line in buf).strip()
        if text:
            blocks.append(Block(kind="para", text=strip_inline(text), section=section_idx))

    para_buf: List[str] = []

    while i < len(lines):
        ln = lines[i]

        # code block fences
        if ln.strip().startswith("```"):
            if in_code:
                blocks.append(Block(kind="code", items=list(code_buf), section=section_idx))
                code_buf = []
                in_code = False
            else:
                if para_buf:
                    flush_paragraph(para_buf); para_buf = []
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(ln)
            i += 1
            continue

        # blank line → flush paragraph
        if ln.strip() == "":
            if para_buf:
                flush_paragraph(para_buf); para_buf = []
            i += 1
            continue

        # horizontal rule
        if re.match(r"^\s*---+\s*$", ln):
            if para_buf:
                flush_paragraph(para_buf); para_buf = []
            blocks.append(Block(kind="hr", section=section_idx))
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", ln)
        if m:
            if para_buf:
                flush_paragraph(para_buf); para_buf = []
            level = len(m.group(1))
            text = strip_inline(m.group(2))
            kind = f"h{level}"
            if level == 2:
                section_idx += 1
            blocks.append(Block(kind=kind, text=text, section=section_idx))
            i += 1
            continue

        # blockquote
        if ln.startswith(">"):
            if para_buf:
                flush_paragraph(para_buf); para_buf = []
            quote_lines: List[str] = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(strip_inline(lines[i].lstrip("> ").rstrip()))
                i += 1
            blocks.append(Block(kind="blockquote", items=quote_lines, section=section_idx))
            continue

        # table: detect a row starting with `|` and the next being a separator
        if ln.lstrip().startswith("|"):
            # Look ahead for separator row
            if i + 1 < len(lines) and re.match(r"^\s*\|?\s*[:\- ]+\|", lines[i + 1]):
                if para_buf:
                    flush_paragraph(para_buf); para_buf = []
                headers = parse_table_row(ln)
                # consume separator
                i += 2
                rows: List[List[str]] = []
                while i < len(lines) and lines[i].lstrip().startswith("|"):
                    rows.append(parse_table_row(lines[i]))
                    i += 1
                # Normalise row widths to header
                width = len(headers)
                rows = [r + [""] * (width - len(r)) if len(r) < width else r[:width] for r in rows]
                blocks.append(Block(kind="table", headers=headers, rows=rows, section=section_idx))
                continue

        # list item
        m = re.match(r"^\s*([-*])\s+(.*)$", ln)
        if m:
            if para_buf:
                flush_paragraph(para_buf); para_buf = []
            items: List[str] = [strip_inline(m.group(2))]
            i += 1
            while i < len(lines):
                m2 = re.match(r"^\s*([-*])\s+(.*)$", lines[i])
                if not m2:
                    break
                items.append(strip_inline(m2.group(2)))
                i += 1
            blocks.append(Block(kind="ul", items=items, section=section_idx))
            continue

        m = re.match(r"^\s*\d+\.\s+(.*)$", ln)
        if m:
            if para_buf:
                flush_paragraph(para_buf); para_buf = []
            items = [strip_inline(m.group(1))]
            i += 1
            while i < len(lines):
                m2 = re.match(r"^\s*\d+\.\s+(.*)$", lines[i])
                if not m2:
                    break
                items.append(strip_inline(m2.group(1)))
                i += 1
            blocks.append(Block(kind="ol", items=items, section=section_idx))
            continue

        # default: accumulate as paragraph
        para_buf.append(ln)
        i += 1

    if para_buf:
        flush_paragraph(para_buf)

    return blocks


# ─── Excel writer ────────────────────────────────────────────────────────────
def _safe_sheet_title(name: str, used: set) -> str:
    bad = set(':\\/?*[]')
    cleaned = "".join(c for c in name if c not in bad).strip()
    cleaned = cleaned[:31] if cleaned else "Sheet"
    base = cleaned
    n = 1
    while cleaned in used:
        n += 1
        suffix = f" {n}"
        cleaned = (base[: 31 - len(suffix)] + suffix).strip()
    used.add(cleaned)
    return cleaned


def write_excel(blocks: List[Block], out_path: str):
    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"

    # Group blocks by H2 section
    sections: List[dict] = []
    current: Optional[dict] = None
    for b in blocks:
        if b.kind == "h1":
            wb_title = b.text
        elif b.kind == "h2":
            current = {"title": b.text, "blocks": []}
            sections.append(current)
        elif current is not None:
            current["blocks"].append(b)

    # ---- Summary sheet
    summary["A1"] = "ERP — Enterprise Admin · Test Flow Document"
    summary["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    summary["A1"].fill = PatternFill("solid", fgColor="1F2937")
    summary["A1"].alignment = Alignment(horizontal="center", vertical="center")
    summary.merge_cells("A1:C1")
    summary.row_dimensions[1].height = 28

    summary["A3"] = "Total sections"
    summary["B3"] = len(sections)
    summary["A4"] = "Tables across all sections"
    summary["B4"] = sum(1 for s in sections for b in s["blocks"] if b.kind == "table")
    summary["A5"] = "Source"
    summary["B5"] = "docs/ERP_Enterprise.md"

    # TOC
    summary["A7"] = "Table of Contents"
    summary["A7"].font = Font(bold=True, size=12)
    summary["A8"] = "#"
    summary["B8"] = "Section"
    summary["C8"] = "Tables"
    for c in ("A8", "B8", "C8"):
        summary[c].font = Font(bold=True, color="FFFFFF")
        summary[c].fill = PatternFill("solid", fgColor="374151")
        summary[c].alignment = Alignment(horizontal="center", vertical="center")

    used_titles = {"Summary"}
    sheet_titles: List[str] = []
    for idx, s in enumerate(sections, 1):
        n_tables = sum(1 for b in s["blocks"] if b.kind == "table")
        # Strip leading numbering like "1. Pre-requisites" → "Pre-requisites"
        clean_title = re.sub(r"^\d+\.\s*", "", s["title"])
        sheet_title = _safe_sheet_title(f"{idx:02d} {clean_title}", used_titles)
        sheet_titles.append(sheet_title)
        row = 8 + idx
        summary.cell(row=row, column=1, value=idx).alignment = Alignment(horizontal="center")
        summary.cell(row=row, column=2, value=s["title"])
        summary.cell(row=row, column=3, value=n_tables).alignment = Alignment(horizontal="center")
        # Hyperlink to sheet
        summary.cell(row=row, column=2).hyperlink = f"#'{sheet_title}'!A1"
        summary.cell(row=row, column=2).font = Font(color="2563EB", underline="single")

    summary.column_dimensions["A"].width = 6
    summary.column_dimensions["B"].width = 60
    summary.column_dimensions["C"].width = 12

    # ---- One sheet per section
    bold_white = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1F2937")
    sub_fill = PatternFill("solid", fgColor="E5E7EB")
    soft_fill = PatternFill("solid", fgColor="F9FAFB")
    border = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center")

    for idx, (s, title) in enumerate(zip(sections, sheet_titles), 1):
        ws = wb.create_sheet(title=title)
        # Title row
        ws["A1"] = s["title"]
        ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
        ws["A1"].fill = header_fill
        ws["A1"].alignment = center
        ws.merge_cells("A1:F1")
        ws.row_dimensions[1].height = 24

        # Back-to-summary link
        ws["G1"] = "← Summary"
        ws["G1"].font = Font(color="2563EB", underline="single")
        ws["G1"].hyperlink = "#'Summary'!A1"
        ws["G1"].alignment = center

        row = 3
        col_widths_seen = {}

        for b in s["blocks"]:
            if b.kind in ("h3", "h4"):
                ws.cell(row=row, column=1, value=b.text).font = Font(bold=True, size=12 if b.kind == "h3" else 11)
                ws.cell(row=row, column=1).fill = sub_fill
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                ws.row_dimensions[row].height = 20
                row += 2
            elif b.kind == "para":
                if not b.text:
                    continue
                ws.cell(row=row, column=1, value=b.text).alignment = wrap
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                row += 2
            elif b.kind == "ul":
                for it in b.items:
                    ws.cell(row=row, column=1, value="• " + it).alignment = wrap
                    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                    row += 1
                row += 1
            elif b.kind == "ol":
                for n, it in enumerate(b.items, 1):
                    ws.cell(row=row, column=1, value=f"{n}. {it}").alignment = wrap
                    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                    row += 1
                row += 1
            elif b.kind == "blockquote":
                for q in b.items:
                    c = ws.cell(row=row, column=1, value="› " + q)
                    c.font = Font(italic=True, color="6B7280")
                    c.alignment = wrap
                    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                    row += 1
                row += 1
            elif b.kind == "code":
                for cl in b.items:
                    c = ws.cell(row=row, column=1, value=cl)
                    c.font = Font(name="Consolas", size=10)
                    c.fill = PatternFill("solid", fgColor="F3F4F6")
                    c.alignment = Alignment(wrap_text=False, vertical="top")
                    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                    row += 1
                row += 1
            elif b.kind == "table":
                # Header row
                for col, h in enumerate(b.headers, 1):
                    c = ws.cell(row=row, column=col, value=h)
                    c.font = bold_white
                    c.fill = header_fill
                    c.alignment = center
                    c.border = border
                ws.row_dimensions[row].height = 22
                # Data rows
                for r_i, r in enumerate(b.rows):
                    rr = row + 1 + r_i
                    for col, v in enumerate(r, 1):
                        c = ws.cell(row=rr, column=col, value=v)
                        c.alignment = wrap
                        c.border = border
                        if r_i % 2 == 0:
                            c.fill = soft_fill
                # Track widths
                widths = []
                for col in range(1, len(b.headers) + 1):
                    max_len = len(str(b.headers[col - 1] or ""))
                    for r in b.rows:
                        v = str(r[col - 1] if col - 1 < len(r) else "")
                        for line in v.split("\n"):
                            if len(line) > max_len:
                                max_len = len(line)
                    width = min(max(max_len + 2, 12), 80)
                    widths.append(width)
                # Apply column widths if greater than current
                for col, w in enumerate(widths, 1):
                    letter = get_column_letter(col)
                    cur = col_widths_seen.get(letter, 0)
                    if w > cur:
                        col_widths_seen[letter] = w
                row += len(b.rows) + 2
            elif b.kind == "hr":
                row += 1

        # Apply max col widths captured for tables
        for letter, w in col_widths_seen.items():
            ws.column_dimensions[letter].width = w
        # Sensible defaults if no tables
        if not col_widths_seen:
            for c in "ABCDEFGH":
                ws.column_dimensions[c].width = 22

        ws.freeze_panes = "A2"

    wb.save(out_path)


# ─── Word writer ─────────────────────────────────────────────────────────────
def _set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D1D5DB")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def write_docx(blocks: List[Block], out_path: str):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for b in blocks:
        if b.kind == "h1":
            p = doc.add_paragraph()
            run = p.add_run(b.text)
            run.bold = True
            run.font.size = Pt(20)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        elif b.kind == "h2":
            doc.add_heading(b.text, level=1)
        elif b.kind == "h3":
            doc.add_heading(b.text, level=2)
        elif b.kind == "h4":
            doc.add_heading(b.text, level=3)
        elif b.kind == "para":
            if b.text:
                doc.add_paragraph(b.text)
        elif b.kind == "ul":
            for it in b.items:
                doc.add_paragraph(it, style="List Bullet")
        elif b.kind == "ol":
            for it in b.items:
                doc.add_paragraph(it, style="List Number")
        elif b.kind == "blockquote":
            for q in b.items:
                p = doc.add_paragraph()
                run = p.add_run(q)
                run.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                p.paragraph_format.left_indent = Cm(0.6)
        elif b.kind == "code":
            for cl in b.items:
                p = doc.add_paragraph()
                run = p.add_run(cl if cl else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        elif b.kind == "table":
            cols = len(b.headers) or 1
            tbl = doc.add_table(rows=1 + len(b.rows), cols=cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit = False
            # Headers
            for c, h in enumerate(b.headers):
                cell = tbl.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_shading(cell, "1F2937")
                _set_cell_borders(cell)
            # Rows
            for r_i, r in enumerate(b.rows):
                for c_i, v in enumerate(r):
                    if c_i >= cols:
                        break
                    cell = tbl.rows[1 + r_i].cells[c_i]
                    cell.text = ""
                    for k, line in enumerate(v.split("\n")):
                        if k == 0:
                            p = cell.paragraphs[0]
                        else:
                            p = cell.add_paragraph()
                        p.add_run(line)
                    _set_cell_borders(cell)
                    if r_i % 2 == 0:
                        _set_cell_shading(cell, "F9FAFB")
        elif b.kind == "hr":
            p = doc.add_paragraph()
            p.add_run("―" * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)


def main():
    blocks = parse_markdown(SRC)
    write_excel(blocks, XLSX)
    write_docx(blocks, DOCX)
    n_h2 = sum(1 for b in blocks if b.kind == "h2")
    n_tab = sum(1 for b in blocks if b.kind == "table")
    print(f"Parsed {len(blocks)} blocks · {n_h2} sections · {n_tab} tables")
    print(f"Wrote: {XLSX}")
    print(f"Wrote: {DOCX}")


if __name__ == "__main__":
    main()
